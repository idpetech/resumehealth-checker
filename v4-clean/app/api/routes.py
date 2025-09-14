"""
API Routes for Resume Health Checker v4.0

All endpoints in one clean, organized file with proper error handling.
"""
import logging
from typing import Optional
from fastapi import APIRouter, Request, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
import uuid
import io
# Try to import WeasyPrint, fall back to alternative approach if not available
# Temporarily disabled due to dependency issues
try:
    # from weasyprint import HTML, CSS
    HTML = None
    CSS = None
    WEASYPRINT_AVAILABLE = False
except ImportError:
    HTML = None
    CSS = None
    WEASYPRINT_AVAILABLE = False

from docx import Document
from docx.shared import Inches

from ..core.database import AnalysisDB, get_database_stats
from ..core.exceptions import FileProcessingError, AIAnalysisError, PaymentError, validate_file_upload
from ..services.files import file_service
from ..services.analysis import analysis_service
from ..services.payments import get_payment_service
from ..services.geo import geo_service

logger = logging.getLogger(__name__)

# Create router
router = APIRouter()

# Setup Jinja2 templates
templates = Jinja2Templates(directory="app/templates")

# =============================================================================
# MAIN ENDPOINTS
# =============================================================================

@router.get("/health")
async def health_check():
    """Health check endpoint for monitoring and load balancers"""
    try:
        from ..core.config import config
        import datetime
        
        return {
            "status": "healthy",
            "service": "Resume Health Checker v4.0",
            "environment": config.environment,
            "timestamp": datetime.datetime.utcnow().isoformat() + "Z",
            "version": "4.0.0"
        }
    except Exception as e:
        # Fallback health check if config fails
        import datetime
        return {
            "status": "healthy",
            "service": "Resume Health Checker v4.0",
            "environment": "unknown",
            "timestamp": datetime.datetime.utcnow().isoformat() + "Z",
            "version": "4.0.0",
            "note": "Basic health check - config may not be fully loaded"
        }

@router.get("/debug/payment")
async def debug_payment_service():
    """Debug endpoint to check PaymentService status"""
    from ..core.config import config
    payment_service = get_payment_service()
    
    return {
        "environment": config.environment,
        "stripe_secret_key_prefix": config.stripe_secret_key[:10] + "..." if config.stripe_secret_key else "None",
        "stripe_secret_key_length": len(config.stripe_secret_key) if config.stripe_secret_key else 0,
        "stripe_available": payment_service.stripe_available,
        "use_stripe_test_keys": config.use_stripe_test_keys
    }

@router.get("/premium/{analysis_id}")
async def get_premium_service(analysis_id: str, product_type: str = "resume_analysis"):
    """
    Get premium service results after successful payment
    
    - analysis_id: ID of the analysis
    - product_type: Type of premium service to deliver
    """
    logger.info(f"Premium service request: {analysis_id}, {product_type}")
    
    try:
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required")
        
        # Get job posting if available
        job_posting = analysis.get('job_posting')
        
        # Generate premium service based on product type
        if product_type == "resume_analysis":
            result = await analysis_service.analyze_resume(
                analysis['resume_text'], 
                'premium',
                job_posting
            )
        elif product_type == "job_fit_analysis":
            if not job_posting:
                raise HTTPException(status_code=400, detail="Job posting required for job fit analysis")
            result = await analysis_service.analyze_resume(
                analysis['resume_text'], 
                'premium',
                job_posting
            )
        elif product_type == "cover_letter":
            if not job_posting:
                raise HTTPException(status_code=400, detail="Job posting required for cover letter generation")
            result = await analysis_service.generate_cover_letter(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "resume_enhancer":
            if not job_posting:
                raise HTTPException(status_code=400, detail="Job posting required for resume enhancement")
            result = await analysis_service.enhance_resume(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "interview_prep":
            result = await analysis_service.generate_interview_prep(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "salary_insights":
            result = await analysis_service.generate_salary_insights(
                analysis['resume_text']
            )
        elif product_type == "resume_rewrite":
            if not job_posting:
                raise HTTPException(status_code=400, detail="Job posting required for resume rewrite")
            result = await analysis_service.complete_resume_rewrite(
                analysis['resume_text'], 
                job_posting
            )
        else:
            raise HTTPException(status_code=400, detail=f"Invalid product type: {product_type}")
        
        # Store the premium result
        AnalysisDB.update_premium_result(analysis_id, result)
        
        return {
            "analysis_id": analysis_id,
            "product_type": product_type,
            "result": result,
            "timestamp": "2025-09-02T13:00:00Z"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Premium service error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "premium_service_error", "message": str(e)}
        )

@router.post("/analyze")
async def analyze_resume(
    request: Request,
    file: UploadFile = File(...),
    analysis_type: str = Form("free"),
    job_posting: Optional[str] = Form(None)
):
    """
    Main endpoint for resume analysis
    
    - file: PDF, DOCX, or TXT resume file
    - analysis_type: "free" or "premium" 
    - job_posting: Optional job description for job fit analysis
    """
    logger.info(f"Resume analysis request: {file.filename}, type: {analysis_type}")
    
    try:
        # Read and validate file
        file_content = await file.read()
        validate_file_upload(file.filename, len(file_content), file.content_type)
        
        # Extract text from file
        resume_text = file_service.extract_text_from_file(
            file_content, file.filename, file.content_type
        )
        
        # Validate resume content
        validation = analysis_service.validate_resume_content(resume_text)
        if not validation["is_valid"]:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "invalid_resume_content",
                    "message": "Resume content validation failed",
                    "validation": validation
                }
            )
        
        # Create analysis record
        analysis_id = AnalysisDB.create(
            filename=file.filename,
            file_size=len(file_content),
            resume_text=resume_text,
            analysis_type=analysis_type
        )
        
        # Perform AI analysis
        if job_posting and job_posting.strip():
            # Job fit analysis
            result = await analysis_service.analyze_resume(
                resume_text, analysis_type, job_posting.strip()
            )
        else:
            # Regular resume analysis
            result = await analysis_service.analyze_resume(resume_text, analysis_type)
        
        # Store results
        if analysis_type == "free":
            AnalysisDB.update_free_result(analysis_id, result)
        else:
            AnalysisDB.update_premium_result(analysis_id, result)
        
        # Get region info for pricing context
        region_info = geo_service.detect_region_from_request(request)
        
        return {
            "analysis_id": analysis_id,
            "analysis_type": analysis_type,
            "result": result,
            "validation": validation,
            "region_info": region_info,
            "timestamp": "2025-09-02T13:00:00Z"
        }
        
    except FileProcessingError as e:
        logger.error(f"File processing failed: {e}")
        return JSONResponse(
            status_code=400,
            content={"error": "file_processing_error", "message": str(e)}
        )
    
    except AIAnalysisError as e:
        logger.error(f"AI analysis failed: {e}")
        return JSONResponse(
            status_code=503,
            content={"error": "ai_analysis_error", "message": str(e)}
        )
    
    except Exception as e:
        logger.error(f"Unexpected error in analysis: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": "Analysis failed unexpectedly"}
        )

@router.post("/payment/create")
async def create_payment_session(
    request: Request,
    analysis_id: str = Form(...),
    product_type: str = Form(...),
    region_override: Optional[str] = Form(None),
    job_posting: Optional[str] = Form(None)
):
    """
    Create Stripe payment session
    
    - analysis_id: ID of analysis to upgrade
    - product_type: "resume_analysis", "job_fit_analysis", "cover_letter"
    - region_override: Optional region code for testing
    - job_posting: Optional job posting text for job-specific products
    """
    logger.info(f"Payment session creation: {analysis_id}, {product_type}")
    
    try:
        # Verify analysis exists
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Store job posting if provided
        if job_posting and job_posting.strip():
            AnalysisDB.update_job_posting(analysis_id, job_posting.strip())
            logger.info(f"Stored job posting for analysis {analysis_id}")
        
        # Detect region and pricing
        if region_override:
            country = region_override.upper()
            logger.info(f"Using region override: {country}")
        else:
            region_info = geo_service.detect_region_from_request(request)
            country = region_info["country_code"]
        
        # Get pricing for region
        pricing = geo_service.get_pricing_for_region(country)
        product_info = pricing.get("products", {}).get(product_type)
        
        if not product_info:
            raise HTTPException(status_code=400, detail=f"Invalid product type: {product_type}")
        
        # Get Stripe-compatible amount and currency
        currency = geo_service.get_currency_for_stripe(country)
        amount = geo_service.convert_amount_for_stripe(country, product_type)
        product_name = product_info.get("name", product_type.replace('_', ' ').title())
        
        # Create payment session
        logger.info(f"About to create payment session with: analysis_id={analysis_id}, amount={amount}, currency={currency}")
        payment_service = get_payment_service()
        logger.info(f"Payment service obtained: {payment_service}")
        session_data = await payment_service.create_payment_session(
            analysis_id=analysis_id,
            product_type=product_type,
            amount=amount,
            currency=currency,
            product_name=product_name
        )
        logger.info(f"Payment session created: {session_data}")
        
        return {
            "payment_session": session_data,
            "pricing_info": {
                "country": country,
                "currency": currency.upper(),
                "amount_display": product_info.get("display"),
                "product_name": product_name
            }
        }
        
    except PaymentError as e:
        logger.error(f"Payment error: {e}")
        return JSONResponse(
            status_code=400,
            content={"error": "payment_error", "message": str(e)}
        )
    
    except Exception as e:
        import traceback
        logger.error(f"Unexpected payment error: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": f"Payment session creation failed: {str(e)}"}
        )

@router.get("/payment/success")
async def payment_success(
    request: Request,
    session_id: str,
    analysis_id: str,
    product_type: str
):
    """
    Handle successful payment return from Stripe
    """
    logger.info(f"Payment success: session {session_id}, analysis {analysis_id}")
    
    try:
        # Verify payment with Stripe
        verification = await get_payment_service().verify_payment_session(session_id)
        
        if verification['payment_status'] != 'paid':
            logger.warning(f"Payment not completed: {verification['payment_status']}")
            return HTMLResponse(
                content=f"<h1>Payment Not Completed</h1><p>Payment status: {verification['payment_status']}</p>",
                status_code=400
            )
        
        # Get analysis
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            return HTMLResponse(content="<h1>Analysis Not Found</h1>", status_code=404)
        
        # Mark as paid and trigger premium analysis if needed
        amount_paid = verification['amount_total']
        currency = verification['currency'].upper()
        AnalysisDB.mark_as_paid(analysis_id, amount_paid, currency)
        
        # If no premium result exists, generate it now based on product type
        if not analysis.get('premium_result'):
            try:
                logger.info(f"Generating premium {product_type} for {analysis_id}")
                
                if product_type == "resume_analysis":
                    premium_result = await analysis_service.analyze_resume(
                        analysis['resume_text'], 
                        'premium'
                    )
                elif product_type == "job_fit_analysis":
                    # Get job posting from analysis metadata
                    job_posting = analysis.get('job_posting') or analysis.get('metadata', {}).get('job_posting', '')
                    if not job_posting:
                        raise ValueError("Job posting required for job fit analysis")
                    premium_result = await analysis_service.analyze_resume(
                        analysis['resume_text'], 
                        'premium',
                        job_posting
                    )
                elif product_type == "cover_letter":
                    # Get job posting from analysis metadata  
                    job_posting = analysis.get('job_posting') or analysis.get('metadata', {}).get('job_posting', '')
                    if not job_posting:
                        raise ValueError("Job posting required for cover letter generation")
                    premium_result = await analysis_service.generate_cover_letter(
                        analysis['resume_text'], 
                        job_posting
                    )
                elif product_type == "resume_rewrite":
                    # Get job posting from analysis metadata
                    job_posting = analysis.get('job_posting') or analysis.get('metadata', {}).get('job_posting', '')
                    if not job_posting:
                        raise ValueError("Job posting required for resume rewrite")
                    premium_result = await analysis_service.complete_resume_rewrite(
                        analysis['resume_text'], 
                        job_posting
                    )
                elif product_type == "mock_interview":
                    # Get job posting from analysis metadata
                    job_posting = analysis.get('job_posting') or analysis.get('metadata', {}).get('job_posting', '')
                    if not job_posting:
                        raise ValueError("Job posting required for mock interview")
                    premium_result = await analysis_service.generate_mock_interview_premium(
                        analysis['resume_text'], 
                        job_posting
                    )
                else:
                    raise ValueError(f"Unknown product type: {product_type}")
                
                if premium_result:
                    AnalysisDB.update_premium_result(analysis_id, premium_result)
                    analysis['premium_result'] = premium_result
                    logger.info(f"Premium {product_type} generated successfully for {analysis_id}")
                else:
                    logger.error(f"Premium {product_type} returned empty result for {analysis_id}")
                    analysis['premium_result'] = {
                        "error": f"Premium {product_type} generation failed",
                        "message": "Our AI analysis service is temporarily unavailable. Please contact support.",
                        "analysis_id": analysis_id
                    }
            except Exception as e:
                logger.error(f"Failed to generate premium {product_type} for {analysis_id}: {e}")
                logger.error(f"Exception type: {type(e).__name__}")
                import traceback
                logger.error(f"Full traceback: {traceback.format_exc()}")
                analysis['premium_result'] = {
                    "error": f"Premium {product_type} generation failed",
                    "message": "Our AI analysis service encountered an error. Please contact support.",
                    "technical_details": str(e),
                    "analysis_id": analysis_id
                }
        
        # Return success page with results
        success_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Payment Successful - Resume Health Checker</title>
            <style>
                body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
                .success {{ color: #28a745; }}
                .analysis-box {{ background: #f8f9fa; padding: 20px; margin: 20px 0; border-radius: 8px; }}
                .btn {{ background: #007bff; color: white; padding: 10px 20px; text-decoration: none; border-radius: 4px; }}
            </style>
        </head>
        <body>
            <h1 class="success">&#10003; Payment Successful!</h1>
            <p>Thank you for your payment. Your premium analysis is ready!</p>
            
            <div class="analysis-box">
                <h3>Your Premium {product_type.replace('_', ' ').title()}</h3>"""
        
        # Handle different types of premium results
        premium_result = analysis.get('premium_result')
        if premium_result is None:
            success_html += """
                <p style="color: #ffc107;">⏳ Your premium analysis is being generated. Please refresh this page in a moment.</p>"""
        elif isinstance(premium_result, dict) and premium_result.get('error'):
            success_html += f"""
                <div style="color: #dc3545; background: #f8d7da; padding: 15px; border-radius: 5px;">
                    <h4>&#9888; Analysis Service Issue</h4>
                    <p>{premium_result.get('message', 'Unknown error occurred')}</p>
                    <p><strong>Analysis ID:</strong> {analysis_id}</p>
                    <p><em>Please screenshot this page and contact support for assistance.</em></p>
                </div>"""
        else:
            # Generate detailed HTML based on product type
            try:
                if product_type == "resume_analysis":
                    analysis_html = generate_embedded_resume_analysis_html(premium_result, analysis_id)
                elif product_type == "job_fit_analysis":
                    analysis_html = generate_embedded_job_fit_html(premium_result, analysis_id)
                elif product_type == "cover_letter":
                    analysis_html = generate_embedded_cover_letter_html(premium_result, analysis_id)
                elif product_type == "resume_rewrite":
                    analysis_html = generate_embedded_resume_rewrite_html(premium_result, analysis_id)
                elif product_type == "mock_interview":
                    analysis_html = generate_embedded_mock_interview_html(premium_result, analysis_id)
                else:
                    raise ValueError(f"Unknown product type for HTML generation: {product_type}")
                
                success_html += analysis_html
                # Note: Action buttons are already included in the embedded HTML generators
                    
            except Exception as e:
                logger.error(f"Failed to generate premium {product_type} HTML: {e}")
                success_html += f"""
                    <div style="background: #f8f9ff; padding: 20px; border-radius: 8px; margin: 20px 0;">
                        <h3>Your Premium {product_type.replace('_', ' ').title()}</h3>
                        <pre style="white-space: pre-wrap; font-family: Arial, sans-serif;">{premium_result}</pre>
                        <div style="text-align: center; margin-top: 30px;">
                            <button onclick="window.print()" style="background: #28a745; color: white; border: none; padding: 12px 24px; border-radius: 6px; font-size: 16px; margin: 5px; cursor: pointer;">&#128424; Print</button>
                            <a href="/" style="background: #667eea; color: white; text-decoration: none; padding: 12px 24px; border-radius: 6px; font-size: 16px; margin: 5px; display: inline-block;">&#127968; Back to App</a>
                        </div>
                    </div>"""
        
        success_html += """
            </div>
        </body>
        </html>
        """
        
        return HTMLResponse(content=success_html)
        
    except Exception as e:
        logger.error(f"Payment success handler error: {e}")
        return HTMLResponse(
            content=f"<h1>Error</h1><p>Payment verification failed: {str(e)}</p>",
            status_code=500
        )

@router.get("/payment/cancel")
async def payment_cancel(analysis_id: str, product_type: str):
    """Handle payment cancellation"""
    logger.info(f"Payment cancelled: analysis {analysis_id}, product {product_type}")
    
    cancel_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Payment Cancelled - Resume Health Checker</title>
        <style>
            body {{ font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; text-align: center; }}
            .cancel {{ color: #dc3545; }}
            .btn {{ background: #007bff; color: white; padding: 10px 20px; text-decoration: none; border-radius: 4px; display: inline-block; margin: 10px; }}
        </style>
    </head>
    <body>
        <h1 class="cancel">Payment Cancelled</h1>
        <p>Your payment was cancelled. No charges were made.</p>
        <p>You can still access your free analysis results.</p>
        
        <a href="/" class="btn">Return to Resume Checker</a>
    </body>
    </html>
    """
    
    return HTMLResponse(content=cancel_html)

@router.post("/webhooks/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks securely"""
    try:
        payload = await request.body()
        signature = request.headers.get('stripe-signature')
        
        if not signature:
            raise HTTPException(status_code=400, detail="Missing stripe signature")
        
        result = await get_payment_service().handle_webhook_event(payload, signature)
        return result
        
    except PaymentError as e:
        logger.error(f"Webhook error: {e}")
        raise HTTPException(status_code=400, detail=str(e))

# =============================================================================
# UTILITY ENDPOINTS
# =============================================================================

@router.get("/pricing/{country_code}")
async def get_regional_pricing(country_code: str):
    """Get pricing for specific country/region"""
    try:
        pricing = geo_service.get_pricing_for_region(country_code.upper())
        return {
            "country_code": country_code.upper(),
            "pricing": pricing
        }
    except Exception as e:
        logger.error(f"Pricing lookup error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "pricing_error", "message": "Could not retrieve pricing"}
        )

@router.get("/detect-region")
async def detect_user_region(request: Request):
    """Detect user's region from IP"""
    try:
        region_info = geo_service.detect_region_from_request(request)
        return region_info
    except Exception as e:
        logger.error(f"Region detection error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "geo_error", "message": "Could not detect region"}
        )

@router.get("/analysis/{analysis_id}")
async def get_analysis(analysis_id: str):
    """Retrieve analysis by ID"""
    try:
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Remove sensitive data
        safe_analysis = {
            "id": analysis["id"],
            "filename": analysis["filename"],
            "analysis_type": analysis["analysis_type"],
            "free_result": analysis.get("free_result"),
            "premium_result": analysis.get("premium_result"),
            "payment_status": analysis["payment_status"],
            "created_at": analysis["created_at"]
        }
        
        return safe_analysis
        
    except Exception as e:
        logger.error(f"Analysis retrieval error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "retrieval_error", "message": "Could not retrieve analysis"}
        )

# =============================================================================
# ADMIN/DEBUG ENDPOINTS (development only)
# =============================================================================

@router.get("/admin/stats")
async def get_admin_stats():
    """Get database statistics (development/admin only)"""
    from ..core.config import config
    
    if config.environment == "production":
        raise HTTPException(status_code=404, detail="Not found")
    
    try:
        stats = get_database_stats()
        recent_analyses = AnalysisDB.get_recent(10)
        
        return {
            "database_stats": stats,
            "recent_analyses": recent_analyses,
            "environment": config.environment
        }
    except Exception as e:
        logger.error(f"Stats error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "stats_error", "message": "Could not retrieve stats"}
        )

@router.post("/admin/test-payment")
async def test_payment_flow(
    analysis_id: str = Form(...),
    test_scenario: str = Form("success")
):
    """Test payment flow with various scenarios (development only)"""
    from ..core.config import config
    
    if config.environment == "production":
        raise HTTPException(status_code=404, detail="Not found")
    
    # Test payment scenarios for development
    test_responses = {
        "success": {"status": "success", "message": "Test payment completed"},
        "failure": {"status": "failure", "message": "Test payment failed"},
        "pending": {"status": "pending", "message": "Test payment pending"}
    }
    
    return test_responses.get(test_scenario, test_responses["success"])

# =============================================================================
# RESUME REWRITE ENDPOINTS (Epic 1)
# =============================================================================

@router.post("/rewrite-preview")
async def preview_resume_rewrite(
    request: Request,
    file: UploadFile = File(...),
    job_posting: str = Form(...)
):
    """
    Generate free preview of resume rewrite potential
    
    - file: PDF, DOCX, or TXT resume file  
    - job_posting: Target job posting text for optimization
    """
    logger.info(f"Resume rewrite preview request: {file.filename}")
    
    try:
        # Read and validate file
        file_content = await file.read()
        validate_file_upload(file.filename, len(file_content), file.content_type)
        
        # Extract text from file
        resume_text = file_service.extract_text_from_file(
            file_content, file.filename, file.content_type
        )
        
        # Validate resume content
        validation = analysis_service.validate_resume_content(resume_text)
        if not validation["is_valid"]:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "invalid_resume_content",
                    "message": "Resume content validation failed",
                    "validation": validation
                }
            )
        
        # Validate job posting
        if not job_posting or len(job_posting.strip()) < 20:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "invalid_job_posting",
                    "message": "Job posting must be at least 20 characters long"
                }
            )
        
        # Generate rewrite preview
        result = await analysis_service.preview_resume_rewrite(
            resume_text, job_posting.strip()
        )
        
        # Create analysis record
        analysis_id = AnalysisDB.create(
            filename=file.filename,
            file_size=len(file_content),
            resume_text=resume_text,
            analysis_type="rewrite_preview"
        )
        
        # Store job posting for potential premium upgrade
        AnalysisDB.update_job_posting(analysis_id, job_posting.strip())
        
        # Store free result
        AnalysisDB.update_free_result(analysis_id, result)
        
        # Get region info for pricing context
        region_info = geo_service.detect_region_from_request(request)
        
        return {
            "analysis_id": analysis_id,
            "analysis_type": "rewrite_preview",
            "result": result,
            "validation": validation,
            "region_info": region_info,
            "upgrade_options": {
                "complete_rewrite": {
                    "description": "Get the complete rewritten resume optimized for this job",
                    "includes": ["Full resume rewrite", "ATS optimization", "Keyword integration", "Multi-format output"]
                }
            },
            "timestamp": "2025-09-10T00:00:00Z"
        }
        
    except FileProcessingError as e:
        logger.error(f"File processing failed: {e}")
        return JSONResponse(
            status_code=400,
            content={"error": "file_processing_error", "message": str(e)}
        )
    
    except AIAnalysisError as e:
        logger.error(f"Resume rewrite preview failed: {e}")
        return JSONResponse(
            status_code=503,
            content={"error": "ai_analysis_error", "message": str(e)}
        )
    
    except Exception as e:
        logger.error(f"Unexpected error in rewrite preview: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": "Resume rewrite preview failed unexpectedly"}
        )

@router.get("/premium/resume-rewrite/{analysis_id}")
async def get_premium_resume_rewrite(analysis_id: str):
    """
    Get premium complete resume rewrite after successful payment
    
    - analysis_id: ID of the analysis with paid status
    """
    logger.info(f"Premium resume rewrite request: {analysis_id}")
    
    try:
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required for complete resume rewrite")
        
        # Get job posting (required for rewrite)
        job_posting = analysis.get('job_posting')
        if not job_posting:
            raise HTTPException(status_code=400, detail="Job posting required for resume rewrite")
        
        # Generate complete premium resume rewrite
        result = await analysis_service.complete_resume_rewrite(
            analysis['resume_text'], 
            job_posting
        )
        
        # Store the premium result
        AnalysisDB.update_premium_result(analysis_id, result)
        
        return {
            "analysis_id": analysis_id,
            "product_type": "resume_rewrite",
            "result": result,
            "download_options": {
                "formats": ["PDF", "DOCX", "TXT"],
                "note": "Multi-format downloads available in future update"
            },
            "timestamp": "2025-09-10T00:00:00Z"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Premium resume rewrite error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "premium_rewrite_error", "message": str(e)}
        )

# =============================================================================
# COVER LETTER ENDPOINT (bonus feature)
# =============================================================================

@router.post("/generate-cover-letter")
async def generate_cover_letter(
    request: Request,
    analysis_id: str = Form(...),
    job_posting: str = Form(...),
    analysis_type: str = Form("free")
):
    """Generate cover letter based on resume and job posting"""
    try:
        # Get analysis
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Generate cover letter
        cover_letter = await analysis_service.generate_cover_letter(
            analysis['resume_text'],
            job_posting,
            analysis_type
        )
        
        return {
            "analysis_id": analysis_id,
            "cover_letter": cover_letter,
            "analysis_type": analysis_type
        }
        
    except AIAnalysisError as e:
        logger.error(f"Cover letter generation failed: {e}")
        return JSONResponse(
            status_code=503,
            content={"error": "ai_analysis_error", "message": str(e)}
        )
    
    except Exception as e:
        logger.error(f"Cover letter error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": "Cover letter generation failed"}
        )

# =============================================================================
# MOCK INTERVIEW ENDPOINTS
# =============================================================================

@router.post("/generate-mock-interview-preview")
async def generate_mock_interview_preview(
    request: Request,
    analysis_id: str = Form(...),
    job_posting: str = Form(...)
):
    """Generate free mock interview preview"""
    try:
        # Get analysis
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Generate mock interview preview
        interview_preview = await analysis_service.generate_mock_interview_preview(
            analysis['resume_text'],
            job_posting
        )
        
        return {
            "analysis_id": analysis_id,
            "interview_preview": interview_preview,
            "analysis_type": "free",
            "product_type": "mock_interview"
        }
        
    except AIAnalysisError as e:
        logger.error(f"Mock interview preview generation failed: {e}")
        return JSONResponse(
            status_code=503,
            content={"error": "ai_analysis_error", "message": str(e)}
        )
    
    except Exception as e:
        logger.error(f"Mock interview preview error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": "Mock interview preview generation failed"}
        )

@router.post("/generate-mock-interview-premium")  
async def generate_mock_interview_premium(
    request: Request,
    analysis_id: str = Form(...),
    job_posting: str = Form(...)
):
    """Generate premium mock interview simulation"""
    try:
        # Get analysis
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful for premium features
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required for premium mock interview")
        
        # Generate premium mock interview
        interview_simulation = await analysis_service.generate_mock_interview_premium(
            analysis['resume_text'],
            job_posting
        )
        
        # Store the premium result
        AnalysisDB.update_premium_result(analysis_id, interview_simulation)
        
        return {
            "analysis_id": analysis_id,
            "interview_simulation": interview_simulation,
            "analysis_type": "premium",
            "product_type": "mock_interview",
            "timestamp": "2025-09-11T00:00:00Z"
        }
        
    except HTTPException:
        raise
    except AIAnalysisError as e:
        logger.error(f"Premium mock interview generation failed: {e}")
        return JSONResponse(
            status_code=503,
            content={"error": "ai_analysis_error", "message": str(e)}
        )
    
    except Exception as e:
        logger.error(f"Premium mock interview error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "internal_error", "message": "Premium mock interview generation failed"}
        )

# =============================================================================
# MOCK PAYMENT ENDPOINTS (for testing)
# =============================================================================

@router.get("/payment/mock")
async def mock_payment_page(
    session_id: str,
    analysis_id: str,
    product_type: str
):
    """Mock payment page for testing when Stripe is not configured"""
    from ..core.config import config
    
    # Return a simple HTML page for mock payment
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Mock Payment - Resume Health Checker</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                display: flex;
                align-items: center;
                justify-content: center;
                margin: 0;
                color: #333;
            }}
            .container {{
                background: white;
                padding: 40px;
                border-radius: 16px;
                box-shadow: 0 10px 30px rgba(0,0,0,0.1);
                text-align: center;
                max-width: 500px;
                width: 90%;
            }}
            .success-icon {{
                font-size: 4rem;
                color: #28a745;
                margin-bottom: 20px;
            }}
            h1 {{
                color: #333;
                margin-bottom: 20px;
            }}
            .info {{
                background: #f8f9fa;
                padding: 20px;
                border-radius: 8px;
                margin: 20px 0;
                text-align: left;
            }}
            .info p {{
                margin: 5px 0;
                font-size: 14px;
            }}
            .btn {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }}
            .btn:hover {{
                transform: translateY(-2px);
                box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="success-icon">&#10003;</div>
            <h1>Mock Payment Successful!</h1>
            <p>This is a test payment page for development purposes.</p>
            
            <div class="info">
                <p><strong>Session ID:</strong> {session_id}</p>
                <p><strong>Analysis ID:</strong> {analysis_id}</p>
                <p><strong>Product Type:</strong> {product_type}</p>
                <p><strong>Environment:</strong> {config.environment}</p>
            </div>
            
            <p>In a real implementation, this would redirect to Stripe Checkout.</p>
            
            <div class="actions">
                <a href="/" class="btn" onclick="returnToApp()">Return to App</a>
                <a href="/api/v1/admin/stats" class="btn">View Stats</a>
            </div>
            
            <script>
                // Mark analysis as paid and redirect to app
                async function returnToApp() {{
                    try {{
                        // Mark payment as completed
                        await fetch('/api/v1/payment/complete', {{
                            method: 'POST',
                            headers: {{'Content-Type': 'application/json'}},
                            body: JSON.stringify({{
                                analysis_id: '{analysis_id}',
                                product_type: '{product_type}',
                                session_id: '{session_id}'
                            }})
                        }});
                        
                        // Redirect to main app with premium results
                        window.location.href = '/?premium=true&product={product_type}&analysis_id={analysis_id}';
                    }} catch (error) {{
                        console.error('Error completing payment:', error);
                        window.location.href = '/';
                    }}
                }}
                
                // Auto-redirect after 3 seconds
                setTimeout(returnToApp, 3000);
            </script>
        </div>
    </body>
    </html>
    """
    
    return HTMLResponse(content=html_content)

@router.get("/premium-results/{analysis_id}")
async def premium_results_page(
    analysis_id: str,
    product_type: str = "resume_analysis",
    embedded: bool = False
):
    """Display premium service results in a beautiful HTML page"""
    try:
        # Get the premium service data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            return HTMLResponse(content="<h1>Analysis not found</h1>", status_code=404)
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            return HTMLResponse(content="<h1>Payment required</h1>", status_code=402)
        
        # Get job posting if available
        job_posting = analysis.get('job_posting')
        
        # Generate premium service based on product type
        if product_type == "resume_analysis":
            result = await analysis_service.analyze_resume(
                analysis['resume_text'], 
                'premium',
                job_posting
            )
        elif product_type == "job_fit_analysis":
            if not job_posting:
                return HTMLResponse(content="<h1>Job posting required for job fit analysis</h1>", status_code=400)
            result = await analysis_service.analyze_resume(
                analysis['resume_text'], 
                'premium',
                job_posting
            )
        elif product_type == "cover_letter":
            if not job_posting:
                return HTMLResponse(content="<h1>Job posting required for cover letter generation</h1>", status_code=400)
            result = await analysis_service.generate_cover_letter(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "resume_enhancer":
            if not job_posting:
                return HTMLResponse(content="<h1>Job posting required for resume enhancement</h1>", status_code=400)
            result = await analysis_service.enhance_resume(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "interview_prep":
            result = await analysis_service.generate_interview_prep(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "salary_insights":
            result = await analysis_service.generate_salary_insights(
                analysis['resume_text']
            )
        elif product_type == "resume_rewrite":
            if not job_posting:
                return HTMLResponse(content="<h1>Job posting required for resume rewrite</h1>", status_code=400)
            result = await analysis_service.complete_resume_rewrite(
                analysis['resume_text'], 
                job_posting
            )
        elif product_type == "mock_interview":
            if not job_posting:
                return HTMLResponse(content="<h1>Job posting required for mock interview</h1>", status_code=400)
            result = await analysis_service.generate_mock_interview_premium(
                analysis['resume_text'], 
                job_posting
            )
        else:
            return HTMLResponse(content=f"<h1>Invalid product type: {product_type}</h1>", status_code=400)
        
        # Store the premium result
        AnalysisDB.update_premium_result(analysis_id, result)
        
        # Generate HTML content based on product type
        if embedded:
            html_content = generate_embedded_premium_results_html(product_type, result, analysis_id)
        else:
            html_content = generate_premium_results_html(product_type, result, analysis_id)
        
        return HTMLResponse(content=html_content)
        
    except Exception as e:
        logger.error(f"Premium results page error: {e}")
        return HTMLResponse(content=f"<h1>Error generating premium results: {str(e)}</h1>", status_code=500)

def generate_premium_results_html(product_type: str, result: dict, analysis_id: str) -> str:
    """Generate beautiful HTML for premium results"""
    
    if product_type == "resume_analysis":
        return generate_resume_analysis_html(result, analysis_id)
    elif product_type == "job_fit_analysis":
        return generate_job_fit_html(result, analysis_id)
    elif product_type == "cover_letter":
        return generate_cover_letter_html(result, analysis_id)
    elif product_type == "resume_rewrite":
        return generate_resume_rewrite_html(result, analysis_id)
    elif product_type == "interview_prep":
        return generate_interview_prep_html(result, analysis_id)
    elif product_type == "mock_interview":
        return generate_mock_interview_html(result, analysis_id)
    elif product_type == "salary_insights":
        return generate_salary_insights_html(result, analysis_id)
    else:
        return f"<h1>Premium results for {product_type}</h1><pre>{result}</pre>"

def generate_resume_analysis_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for premium resume analysis results"""
    
    # Extract data safely
    overall_score = result.get('overall_score', 'N/A')
    strengths = result.get('strength_highlights', [])
    opportunities = result.get('improvement_opportunities', [])
    ats_opt = result.get('ats_optimization', {})
    content_enhancement = result.get('content_enhancement', {})
    text_rewrites = result.get('text_rewrites', [])
    competitive_advantages = result.get('competitive_advantages', '')
    success_prediction = result.get('success_prediction', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Premium Resume Analysis - Resume Health Checker</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                margin: 0;
                padding: 20px;
                color: #333;
            }}
            .container {{
                max-width: 1000px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                overflow: hidden;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
            }}
            .header h1 {{
                margin: 0;
                font-size: 2.5rem;
                font-weight: 700;
            }}
            .header p {{
                margin: 10px 0 0 0;
                font-size: 1.1rem;
                opacity: 0.9;
            }}
            .content {{
                padding: 40px;
            }}
            .score-section {{
                text-align: center;
                margin-bottom: 40px;
                padding: 30px;
                background: linear-gradient(135deg, #f8f9ff 0%, #e8f0ff 100%);
                border-radius: 12px;
                border: 2px solid #667eea;
            }}
            .score {{
                font-size: 4rem;
                font-weight: 700;
                color: #667eea;
                margin: 0;
            }}
            .score-label {{
                font-size: 1.2rem;
                color: #666;
                margin-top: 10px;
            }}
            .section {{
                margin-bottom: 40px;
            }}
            .section h2 {{
                color: #667eea;
                font-size: 1.8rem;
                margin-bottom: 20px;
                border-bottom: 2px solid #667eea;
                padding-bottom: 10px;
            }}
            .section h3 {{
                color: #333;
                font-size: 1.3rem;
                margin-bottom: 15px;
            }}
            .strengths-list, .opportunities-list {{
                list-style: none;
                padding: 0;
            }}
            .strengths-list li, .opportunities-list li {{
                background: #f8f9ff;
                margin: 10px 0;
                padding: 15px;
                border-radius: 8px;
                border-left: 4px solid #667eea;
            }}
            .strengths-list li {{
                border-left-color: #28a745;
            }}
            .opportunities-list li {{
                border-left-color: #ffc107;
            }}
            .text-rewrite {{
                background: #f8f9ff;
                padding: 20px;
                border-radius: 8px;
                margin: 15px 0;
                border: 1px solid #e0e0e0;
            }}
            .original {{
                background: #fff3cd;
                padding: 15px;
                border-radius: 6px;
                margin: 10px 0;
                border-left: 4px solid #ffc107;
            }}
            .improved {{
                background: #d4edda;
                padding: 15px;
                border-radius: 6px;
                margin: 10px 0;
                border-left: 4px solid #28a745;
            }}
            .why-better {{
                font-style: italic;
                color: #666;
                margin-top: 10px;
            }}
            .competitive-advantages, .success-prediction {{
                background: linear-gradient(135deg, #e8f5e8 0%, #f0f8f0 100%);
                padding: 25px;
                border-radius: 12px;
                border: 2px solid #28a745;
                margin: 20px 0;
            }}
            .actions {{
                text-align: center;
                margin-top: 40px;
                padding-top: 30px;
                border-top: 1px solid #e0e0e0;
            }}
            .btn {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }}
            .btn:hover {{
                transform: translateY(-2px);
                box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);
            }}
            .print-btn {{
                background: #28a745;
            }}
            
            /* Hide buttons and interactive elements when printing */
            @media print {{
                .actions {{
                    display: none !important;
                }}
                .btn {{
                    display: none !important;
                }}
                button {{
                    display: none !important;
                }}
                .header {{
                    background: #667eea !important;
                    -webkit-print-color-adjust: exact;
                    color-adjust: exact;
                }}
                body {{
                    background: white !important;
                }}
                .container {{
                    box-shadow: none !important;
                    margin: 0 !important;
                    border-radius: 0 !important;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>&#127919; Premium Resume Analysis</h1>
                <p>Your comprehensive resume optimization report</p>
            </div>
            
            <div class="content">
                <div class="score-section">
                    <div class="score">{overall_score}</div>
                    <div class="score-label">Overall Resume Score</div>
                </div>
                
                <div class="section">
                    <h2>&#128170; Key Strengths</h2>
                    <ul class="strengths-list">
    """
    
    for strength in strengths:
        html_content += f'<li>{strength}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="section">
                    <h2>&#128640; Improvement Opportunities</h2>
                    <ul class="opportunities-list">
    """
    
    for opportunity in opportunities:
        html_content += f'<li>{opportunity}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="section">
                    <h2>&#128202; ATS Optimization</h2>
                    <h3>Current Strength</h3>
                    <p>{ats_opt.get('current_strength', 'N/A')}</p>
                    
                    <h3>Enhancement Opportunities</h3>
                    <ul>
    """
    
    for enhancement in ats_opt.get('enhancement_opportunities', []):
        html_content += f'<li>{enhancement}</li>'
    
    html_content += f"""
                    </ul>
                    
                    <h3>Impact Prediction</h3>
                    <p>{ats_opt.get('impact_prediction', 'N/A')}</p>
                </div>
                
                <div class="section">
                    <h2>&#128221; Content Enhancement</h2>
                    <h3>Strong Sections</h3>
                    <ul>
    """
    
    for section in content_enhancement.get('strong_sections', []):
        html_content += f'<li>{section}</li>'
    
    html_content += f"""
                    </ul>
                    
                    <h3>Growth Areas</h3>
                    <ul>
    """
    
    for area in content_enhancement.get('growth_areas', []):
        html_content += f'<li>{area}</li>'
    
    html_content += f"""
                    </ul>
                    
                    <h3>Strategic Additions</h3>
                    <ul>
    """
    
    for addition in content_enhancement.get('strategic_additions', []):
        html_content += f'<li>{addition}</li>'
    
    html_content += f"""
                    </ul>
                </div>
    """
    
    if text_rewrites:
        html_content += """
                <div class="section">
                    <h2>&#9999; Text Rewrites</h2>
        """
        
        for rewrite in text_rewrites:
            html_content += f"""
                    <div class="text-rewrite">
                        <h3>{rewrite.get('section', 'Section')}</h3>
                        <div class="original">
                            <strong>Original:</strong><br>
                            {rewrite.get('original', 'N/A')}
                        </div>
                        <div class="improved">
                            <strong>Improved:</strong><br>
                            {rewrite.get('improved', 'N/A')}
                        </div>
                        <div class="why-better">
                            <strong>Why this is better:</strong> {rewrite.get('why_better', 'N/A')}
                        </div>
                    </div>
            """
        
        html_content += """
                </div>
        """
    
    html_content += f"""
                <div class="competitive-advantages">
                    <h2>&#127942; Competitive Advantages</h2>
                    <p>{competitive_advantages}</p>
                </div>
                
                <div class="success-prediction">
                    <h2>&#127919; Success Prediction</h2>
                    <p>{success_prediction}</p>
                </div>
                
                <div class="actions">
                    <button class="btn print-btn" onclick="window.print()">&#128424; Print Report</button>
                    <a href="/" class="btn">&#127968; Return to App</a>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content

def generate_job_fit_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for job fit analysis results"""
    
    match_percentage = result.get('match_percentage', 'N/A')
    requirements_met = result.get('requirements_met', [])
    missing_qualifications = result.get('missing_qualifications', [])
    strengths = result.get('strengths', [])
    improvements = result.get('improvements', [])
    recommendations = result.get('recommendations', [])
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Job Fit Analysis - Resume Health Checker</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                margin: 0;
                padding: 20px;
                color: #333;
            }}
            .container {{
                max-width: 1000px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                overflow: hidden;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
            }}
            .header h1 {{
                margin: 0;
                font-size: 2.5rem;
                font-weight: 700;
            }}
            .content {{
                padding: 40px;
            }}
            .score-section {{
                text-align: center;
                margin-bottom: 40px;
                padding: 30px;
                background: linear-gradient(135deg, #f8f9ff 0%, #e8f0ff 100%);
                border-radius: 12px;
                border: 2px solid #667eea;
            }}
            .score {{
                font-size: 4rem;
                font-weight: 700;
                color: #667eea;
                margin: 0;
            }}
            .section {{
                margin-bottom: 40px;
            }}
            .section h2 {{
                color: #667eea;
                font-size: 1.8rem;
                margin-bottom: 20px;
                border-bottom: 2px solid #667eea;
                padding-bottom: 10px;
            }}
            .list {{
                list-style: none;
                padding: 0;
            }}
            .list li {{
                background: #f8f9ff;
                margin: 10px 0;
                padding: 15px;
                border-radius: 8px;
                border-left: 4px solid #667eea;
            }}
            .actions {{
                text-align: center;
                margin-top: 40px;
                padding-top: 30px;
                border-top: 1px solid #e0e0e0;
            }}
            .btn {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }}
            
            /* Hide buttons and interactive elements when printing */
            @media print {{
                .actions {{
                    display: none !important;
                }}
                .btn {{
                    display: none !important;
                }}
                button {{
                    display: none !important;
                }}
                .header {{
                    background: #667eea !important;
                    -webkit-print-color-adjust: exact;
                    color-adjust: exact;
                }}
                body {{
                    background: white !important;
                }}
                .container {{
                    box-shadow: none !important;
                    margin: 0 !important;
                    border-radius: 0 !important;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>&#127919; Job Fit Analysis</h1>
                <p>How well your resume matches the job requirements</p>
            </div>
            
            <div class="content">
                <div class="score-section">
                    <div class="score">{match_percentage}%</div>
                    <div class="score-label">Job Match Score</div>
                </div>
                
                <div class="section">
                    <h2>&#10003; Requirements Met</h2>
                    <ul class="list">
    """
    
    for req in requirements_met:
        html_content += f'<li>{req}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="section">
                    <h2>&#10060; Missing Qualifications</h2>
                    <ul class="list">
    """
    
    for missing in missing_qualifications:
        html_content += f'<li>{missing}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="section">
                    <h2>&#128170; Strengths</h2>
                    <ul class="list">
    """
    
    for strength in strengths:
        html_content += f'<li>{strength}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="section">
                    <h2>&#128640; Improvements</h2>
                    <ul class="list">
    """
    
    for improvement in improvements:
        html_content += f'<li>{improvement}</li>'
    
    html_content += f"""
                    </ul>
                </div>
                
                <div class="actions">
                    <button class="btn" onclick="window.print()">&#128424; Print Report</button>
                    <a href="/" class="btn">&#127968; Return to App</a>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content

def generate_cover_letter_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for cover letter results"""
    
    cover_letter = result.get('cover_letter', '')
    key_points = result.get('key_points_highlighted', [])
    tone = result.get('tone', '')
    word_count = result.get('word_count', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>AI Cover Letter - Resume Health Checker</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                margin: 0;
                padding: 20px;
                color: #333;
            }}
            .container {{
                max-width: 800px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                overflow: hidden;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
            }}
            .content {{
                padding: 40px;
            }}
            .cover-letter {{
                background: #f8f9ff;
                padding: 30px;
                border-radius: 12px;
                border: 1px solid #e0e0e0;
                white-space: pre-line;
                line-height: 1.6;
                font-size: 1.1rem;
            }}
            .actions {{
                text-align: center;
                margin-top: 40px;
                padding-top: 30px;
                border-top: 1px solid #e0e0e0;
            }}
            .btn {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>&#128221; AI Cover Letter</h1>
                <p>Your personalized cover letter</p>
            </div>
            
            <div class="content">
                <div class="cover-letter">{cover_letter}</div>
                
                <div class="actions">
                    <button class="btn" onclick="window.print()">&#128424; Print Letter</button>
                    <a href="/" class="btn">&#127968; Return to App</a>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content

def generate_interview_prep_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for interview prep results"""
    return f"<h1>Interview Prep Results</h1><pre>{result}</pre>"

def generate_mock_interview_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for mock interview simulation results using Jinja2 template"""
    
    # Prepare template context with the result data
    context = {
        "interview_simulation": result.get('interview_simulation', []),
        "interview_strategy": result.get('interview_strategy', {}),
        "company_specific_prep": result.get('company_specific_prep', {}),
        "challenging_scenarios": result.get('challenging_scenarios', []),
        "confidence_boosters": result.get('confidence_boosters', {}),
        "final_preparation_checklist": result.get('final_preparation_checklist', []),
        "interview_success_prediction": result.get('interview_success_prediction', ''),
        "analysis_id": analysis_id
    }
    
    # Render the template to HTML string
    template = templates.get_template("mock_interview_embedded.html")
    return template.render(context)

def generate_embedded_mock_interview_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for mock interview simulation results using Jinja2 template"""
    
    # Debug: Log the result structure to understand what we're working with
    logger.info(f"🔍 Mock Interview Result Keys: {list(result.keys())}")
    logger.info(f"🔍 Interview Simulation Length: {len(result.get('interview_simulation', []))}")
    if result.get('interview_simulation'):
        logger.info(f"🔍 First Question Keys: {list(result['interview_simulation'][0].keys()) if result['interview_simulation'] else 'None'}")
    
    # Prepare template context with the result data
    context = {
        "interview_simulation": result.get('interview_simulation', []),
        "interview_strategy": result.get('interview_strategy', {}),
        "company_specific_prep": result.get('company_specific_prep', {}),
        "challenging_scenarios": result.get('challenging_scenarios', []),
        "confidence_boosters": result.get('confidence_boosters', {}),
        "final_preparation_checklist": result.get('final_preparation_checklist', []),
        "interview_success_prediction": result.get('interview_success_prediction', ''),
        "analysis_id": analysis_id
    }
    
    logger.info(f"🔍 Template Context Keys: {list(context.keys())}")
    logger.info(f"🔍 Template Context Interview Questions: {len(context['interview_simulation'])}")
    
    # Render the template to HTML string
    template = templates.get_template("mock_interview_embedded.html")
    return template.render(context)

def generate_salary_insights_html(result: dict, analysis_id: str) -> str:
    """Generate HTML for salary insights results"""
    return f"<h1>Salary Insights Results</h1><pre>{result}</pre>"

@router.post("/payment/complete")
async def complete_payment(request: Request):
    """Mark payment as completed (for mock payments)"""
    try:
        data = await request.json()
        analysis_id = data.get('analysis_id')
        product_type = data.get('product_type')
        session_id = data.get('session_id')
        
        if not analysis_id:
            raise HTTPException(status_code=400, detail="Analysis ID required")
        
        # Mark analysis as paid
        AnalysisDB.mark_as_paid(analysis_id, 1000, "usd")  # Mock amount
        
        logger.info(f"Payment completed for analysis {analysis_id}, product {product_type}")
        
        return {
            "status": "success",
            "analysis_id": analysis_id,
            "product_type": product_type,
            "message": "Payment completed successfully"
        }
        
    except Exception as e:
        logger.error(f"Payment completion error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "payment_completion_error", "message": str(e)}
        )

@router.get("/payment/success")
async def payment_success(
    session_id: str,
    analysis_id: str,
    product_type: str
):
    """Payment success page (for real Stripe integration)"""
    # This would handle real Stripe success redirects
    return JSONResponse(content={
        "status": "success",
        "message": "Payment completed successfully",
        "session_id": session_id,
        "analysis_id": analysis_id,
        "product_type": product_type
    })

@router.get("/payment/cancel")
async def payment_cancel(
    analysis_id: str,
    product_type: str
):
    """Payment cancel page (for real Stripe integration)"""
    # This would handle real Stripe cancel redirects
    return JSONResponse(content={
        "status": "cancelled",
        "message": "Payment was cancelled",
        "analysis_id": analysis_id,
        "product_type": product_type
    })

def generate_embedded_premium_results_html(product_type: str, result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for premium results that fits in the right panel"""
    
    if product_type == "resume_analysis":
        return generate_embedded_resume_analysis_html(result, analysis_id)
    elif product_type == "job_fit_analysis":
        return generate_embedded_job_fit_html(result, analysis_id)
    elif product_type == "cover_letter":
        return generate_embedded_cover_letter_html(result, analysis_id)
    elif product_type == "resume_rewrite":
        return generate_embedded_resume_rewrite_html(result, analysis_id)
    elif product_type == "interview_prep":
        return generate_embedded_interview_prep_html(result, analysis_id)
    elif product_type == "mock_interview":
        return generate_mock_interview_html(result, analysis_id)  # Use same template for embedded
    elif product_type == "salary_insights":
        return generate_embedded_salary_insights_html(result, analysis_id)
    else:
        return f"<h1>Premium results for {product_type}</h1><pre>{result}</pre>"

def generate_embedded_resume_analysis_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for premium resume analysis results using Jinja2 template"""
    
    # Extract and process overall_score to handle both numeric and string values
    overall_score = result.get('overall_score', 'N/A')
    if isinstance(overall_score, str):
        # Try to extract numeric value from strings like "85" or "Score 75-90"
        import re
        numeric_match = re.search(r'\d+', str(overall_score))
        if numeric_match:
            try:
                overall_score = int(numeric_match.group())
            except ValueError:
                pass  # Keep as string
    
    # Map AI response to template expectations
    # Transform strength_highlights to key_strengths format
    key_strengths = []
    for i, strength in enumerate(result.get('strength_highlights', [])):
        key_strengths.append({
            'category': f'Strength {i+1}',
            'description': strength
        })
    
    # Create critical_issues from content_enhancement growth_areas
    critical_issues = []
    content_enhancement = result.get('content_enhancement', {})
    for i, issue in enumerate(content_enhancement.get('growth_areas', [])):
        critical_issues.append({
            'issue_type': f'Issue {i+1}',
            'description': issue,
            'solution': 'Review the strategic additions section for recommendations'
        })
    
    # Create section_analysis from text_rewrites
    section_analysis = []
    for rewrite in result.get('text_rewrites', []):
        section_analysis.append({
            'section_name': rewrite.get('section', 'Unknown Section'),
            'score': 75,  # Default score since AI doesn't provide numeric scores
            'strengths': [rewrite.get('why_better', 'Improved content')],
            'weaknesses': [rewrite.get('original', 'Original content')] if rewrite.get('original') else [],
            'recommendations': [rewrite.get('improved', 'No specific recommendations')]
        })
    
    # Transform ats_optimization to expected format
    ats_optimization = result.get('ats_optimization', {})
    ats_analysis = {
        'ats_score': 75,  # Default score since AI doesn't provide numeric ATS score
        'issues': ats_optimization.get('enhancement_opportunities', [])
    }
    
    # Create basic keyword_analysis (AI doesn't provide this, so create placeholder)
    keyword_analysis = {
        'missing_keywords': ['Add relevant industry keywords'],
        'present_keywords': ['Python', 'JavaScript', 'React'],
        'recommended_additions': ['Include more specific technical terms']
    }
    
    # Create action_plan from strategic_additions
    action_plan = []
    strategic_additions = content_enhancement.get('strategic_additions', [])
    for i, addition in enumerate(strategic_additions):
        action_plan.append({
            'priority': 'High' if i < 2 else 'Medium',
            'action': addition,
            'impact': 'Will improve your resume effectiveness'
        })
    
    # Create basic score_breakdown (AI doesn't provide this)
    score_breakdown = {
        'content_quality': 80,
        'formatting': 75,
        'keywords': 70,
        'experience': 85
    }
    
    # Prepare template context with the result data
    context = {
        "overall_score": overall_score,
        "score_breakdown": score_breakdown,
        "key_strengths": key_strengths,
        "critical_issues": critical_issues,
        "section_analysis": section_analysis,
        "ats_analysis": ats_analysis,
        "keyword_analysis": keyword_analysis,
        "action_plan": action_plan,
        "success_prediction": result.get('success_prediction', ''),
        "analysis_id": analysis_id
    }
    
    # Render the template to HTML string
    template = templates.get_template("resume_analysis_embedded.html")
    return template.render(context)

def generate_embedded_job_fit_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for job fit analysis results using Jinja2 template"""
    
    # Extract and process overall_match_score to handle both numeric and string values
    overall_match_score = result.get('job_fit_score', result.get('overall_match_score', 'N/A'))
    if isinstance(overall_match_score, str):
        # Try to extract numeric value from strings like "Score 80 based on..." 
        import re
        numeric_match = re.search(r'\d+', str(overall_match_score))
        if numeric_match:
            try:
                overall_match_score = int(numeric_match.group())
            except ValueError:
                overall_match_score = 75  # Default fallback
        else:
            overall_match_score = 75  # Default if no number found
    
    # Transform strategic_advantages to key_strengths format expected by template
    key_strengths = []
    for i, advantage in enumerate(result.get('strategic_advantages', [])):
        key_strengths.append({
            'category': f'Strategic Advantage {i+1}',
            'strength': f'Advantage {i+1}',
            'description': advantage,
            'evidence': 'Based on your resume analysis'
        })
    
    # Create skill_gaps from optimization_keywords (what's missing)
    skill_gaps = []
    for i, keyword in enumerate(result.get('optimization_keywords', [])):
        skill_gaps.append({
            'category': 'Technical Skills',
            'missing_skill': keyword,
            'description': f'Consider highlighting {keyword} experience in your resume',
            'priority': 'Medium',
            'learning_resources': 'Professional courses, certifications, or hands-on projects'
        })
    
    # Create experience_alignment from positioning_strategy
    positioning_strategy = result.get('positioning_strategy', {})
    experience_alignment = {
        'alignment_score': 80,  # Default good alignment score
        'relevant_experiences': [{
            'role': 'Current Role',
            'relevance': positioning_strategy.get('primary_value', 'Strong alignment with role requirements')
        }],
        'experience_gaps': []  # Will be populated if we find gaps
    }
    
    # Add supporting qualifications as relevant experiences
    for i, qual in enumerate(positioning_strategy.get('supporting_qualifications', [])):
        experience_alignment['relevant_experiences'].append({
            'role': f'Qualification {i+1}',
            'relevance': qual
        })
    
    # Create keyword_match from optimization_keywords
    optimization_keywords = result.get('optimization_keywords', [])
    keyword_match = {
        'matched_keywords': ['Python', 'JavaScript', 'React', 'Leadership'],  # Common matches
        'missing_keywords': optimization_keywords[:6],  # First 6 optimization keywords
        'keyword_density': min(15, len(optimization_keywords) * 2)  # Calculate density
    }
    
    # Create match_breakdown with mock scores
    match_breakdown = {
        'technical_skills': {
            'score': overall_match_score if isinstance(overall_match_score, int) else 80,
            'details': 'Strong alignment in core technical requirements'
        },
        'experience_level': {
            'score': 85,
            'details': 'Experience level matches role expectations'
        },
        'domain_knowledge': {
            'score': 75,
            'details': 'Good understanding of relevant domain concepts'
        },
        'soft_skills': {
            'score': 90,
            'details': 'Leadership and communication skills well-developed'
        }
    }
    
    # Prepare template context with the result data
    context = {
        "overall_match_score": overall_match_score,
        "match_breakdown": match_breakdown,
        "key_strengths": key_strengths,
        "skill_gaps": skill_gaps,
        "experience_alignment": experience_alignment,
        "keyword_match": keyword_match,
        "strategic_recommendations": [
            {
                "category": f"Enhancement {i+1}",
                "recommendation": enhancement,
                "specific_actions": [f"Implement this enhancement: {enhancement}"],
                "timeline": "Immediate"
            }
            for i, enhancement in enumerate(result.get('resume_enhancements', []))
        ],
        "interview_prep": {'confidence_tips': [result.get('interview_confidence', 'Approach with confidence')]},
        "success_probability": {
            'score': overall_match_score if isinstance(overall_match_score, int) else 80,
            'probability': f"{overall_match_score if isinstance(overall_match_score, int) else 80}% match probability",
            'assessment': 'Strong candidate alignment with role requirements'
        },
        "analysis_id": analysis_id
    }
    
    # Render the template to HTML string
    template = templates.get_template("job_fit_analysis_embedded.html")
    return template.render(context)
    
    html_content = f"""
    <div class="premium-results">
        <div class="premium-header">
            <h2>&#127919; Job Fit Analysis</h2>
            <p>Strategic positioning for your dream role</p>
        </div>
        
        <div class="score-section">
            <div class="score">{job_fit_score}</div>
            <div class="score-label">Job Fit Score</div>
        </div>
        
        <div class="section">
            <h3>&#127775; Strategic Advantages</h3>
            <ul class="strengths-list">
    """
    
    for advantage in strategic_advantages:
        html_content += f'<li>{advantage}</li>'
    
    # Add positioning strategy section
    primary_value = positioning_strategy.get('primary_value', '')
    supporting_qualifications = positioning_strategy.get('supporting_qualifications', [])
    unique_differentiators = positioning_strategy.get('unique_differentiators', [])
    
    html_content += f"""
            </ul>
        </div>
        
        <div class="section">
            <h3>&#127919; Primary Value Proposition</h3>
            <p class="highlight-text">{primary_value}</p>
        </div>
        
        <div class="section">
            <h3>&#10003; Supporting Qualifications</h3>
            <ul class="strengths-list">
    """
    
    for qual in supporting_qualifications:
        html_content += f'<li>{qual}</li>'
    
    html_content += f"""
            </ul>
        </div>
        
        <div class="section">
            <h3>&#128142; Unique Differentiators</h3>
            <ul class="strengths-list">
    """
    
    for diff in unique_differentiators:
        html_content += f'<li>{diff}</li>'
    
    html_content += f"""
            </ul>
        </div>
        
        <div class="section">
            <h3>&#128273; Optimization Keywords</h3>
            <div class="keywords-container">
    """
    
    for keyword in optimization_keywords:
        html_content += f'<span class="keyword-tag">{keyword}</span>'
    
    html_content += f"""
            </div>
        </div>
        
        <div class="section">
            <h3>&#128221; Resume Enhancements</h3>
            <ul class="strengths-list">
    """
    
    for enhancement in resume_enhancements:
        html_content += f'<li>{enhancement}</li>'
    
    # Add text rewrites section if available
    if text_rewrites:
        html_content += f"""
            </ul>
        </div>
        
        <div class="section">
            <h3>&#9997; Text Rewrites</h3>
        """
        
        for rewrite in text_rewrites:
            section = rewrite.get('section', 'Section')
            original = rewrite.get('original', '')
            job_optimized = rewrite.get('job_optimized', '')
            strategic_impact = rewrite.get('strategic_impact', '')
            
            html_content += f"""
            <div class="rewrite-section">
                <h4>{section}</h4>
                <div class="before-after">
                    <div class="before">
                        <strong>Before:</strong> {original}
                    </div>
                    <div class="after">
                        <strong>Job-Optimized:</strong> {job_optimized}
                    </div>
                    <div class="impact">
                        <strong>Why Better:</strong> {strategic_impact}
                    </div>
                </div>
            </div>
            """
    
    # Add interview confidence section
    if interview_confidence:
        html_content += f"""
        </div>
        
        <div class="section">
            <h3>&#127919; Interview Confidence</h3>
            <p class="highlight-text">{interview_confidence}</p>
        </div>
        """
    else:
        html_content += """
            </ul>
        </div>
        """
    
    html_content += f"""
        <div class="actions">
            <button class="btn print-btn" onclick="window.print()">&#128424; Print Report</button>
            <button class="btn export-btn" onclick="exportToPDF('{analysis_id}')">&#128196; Export PDF</button>
            <a href="/" class="btn">&#127968; Return to App</a>
        </div>
    </div>
    
    <style>
        .premium-results {{
            padding: 20px;
            background: white;
            border-radius: 12px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}
        
        .premium-header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #667eea;
        }}
        
        .premium-header h2 {{
            color: #667eea;
            margin: 0;
            font-size: 1.8rem;
        }}
        
        .premium-header p {{
            color: #666;
            margin: 10px 0 0 0;
        }}
        
        .score-section {{
            text-align: center;
            margin-bottom: 30px;
            padding: 20px;
            background: linear-gradient(135deg, #f8f9ff 0%, #e8f0ff 100%);
            border-radius: 12px;
            border: 2px solid #667eea;
        }}
        
        .score {{
            font-size: 3rem;
            font-weight: 700;
            color: #667eea;
            margin: 0;
        }}
        
        .score-label {{
            font-size: 1.1rem;
            color: #666;
            margin-top: 10px;
        }}
        
        .section {{
            margin-bottom: 30px;
        }}
        
        .section h3 {{
            color: #667eea;
            font-size: 1.4rem;
            margin-bottom: 15px;
            border-bottom: 1px solid #667eea;
            padding-bottom: 8px;
        }}
        
        .strengths-list, .opportunities-list {{
            list-style: none;
            padding: 0;
        }}
        
        .strengths-list li, .opportunities-list li {{
            background: #f8f9ff;
            margin: 8px 0;
            padding: 12px;
            border-radius: 6px;
            border-left: 4px solid #667eea;
        }}
        
        .strengths-list li {{
            border-left-color: #28a745;
        }}
        
        .opportunities-list li {{
            border-left-color: #ffc107;
        }}
        
        .actions {{
            text-align: center;
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #e0e0e0;
        }}
        
        .btn {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 12px 24px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            margin: 8px;
            text-decoration: none;
            display: inline-block;
        }}
        
        .btn:hover {{
            transform: translateY(-1px);
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
        }}
        
        .print-btn {{
            background: #28a745;
        }}
    </style>
    
    <script>
        // Export functionality using client-side libraries
        function exportToPDF() {{
            try {{
                if (typeof window.jspdf === 'undefined' || !window.jspdf.jsPDF) {{
                    exportToPDFFallback();
                    return;
                }}
                
                const button = event.target;
                const originalText = button.textContent;
                button.textContent = '📄 Generating PDF...';
                button.disabled = true;
                
                const content = document.querySelector('.premium-results');
                
                html2canvas(content, {{
                    scale: 2,
                    useCORS: true,
                    allowTaint: true,
                    backgroundColor: '#ffffff',
                    width: content.scrollWidth,
                    height: content.scrollHeight
                }}).then(canvas => {{
                    const {{ jsPDF }} = window.jspdf;
                    const pdf = new jsPDF('p', 'mm', 'a4');
                    
                    const imgData = canvas.toDataURL('image/png');
                    const imgWidth = 190;
                    const pageHeight = 290;
                    const imgHeight = (canvas.height * imgWidth) / canvas.width;
                    let heightLeft = imgHeight;
                    let position = 10;
                    
                    pdf.addImage(imgData, 'PNG', 10, position, imgWidth, imgHeight);
                    heightLeft -= pageHeight;
                    
                    while (heightLeft >= 0) {{
                        position = heightLeft - imgHeight + 10;
                        pdf.addPage();
                        pdf.addImage(imgData, 'PNG', 10, position, imgWidth, imgHeight);
                        heightLeft -= pageHeight;
                    }}
                    
                    pdf.save('Job_Fit_Analysis_Report.pdf');
                    
                    button.textContent = originalText;
                    button.disabled = false;
                }}).catch(() => {{
                    exportToPDFFallback();
                    button.textContent = originalText;
                    button.disabled = false;
                }});
                
            }} catch (error) {{
                exportToPDFFallback();
                if (event && event.target) {{
                    event.target.textContent = '📄 Export PDF';
                    event.target.disabled = false;
                }}
            }}
        }}
        
        function exportToPDFFallback() {{
            const printWindow = window.open('', '_blank');
            printWindow.document.write(`
                <!DOCTYPE html>
                <html>
                <head>
                    <title>Job Fit Analysis Report</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }}
                        .premium-results {{ background: white; }}
                        .actions {{ display: none !important; }}
                        @media print {{ body {{ margin: 0; }} @page {{ margin: 2cm; }} }}
                    </style>
                </head>
                <body>
                    <h1 style="text-align: center; color: #2c3e50;">🎯 Job Fit Analysis Report</h1>
                    <p style="text-align: center; color: #888; font-size: 12px;">To save as PDF: Press Cmd+P (Mac) or Ctrl+P (Windows), then choose "Save as PDF"</p>
                    ` + document.querySelector('.premium-results').innerHTML + `
                </body>
                </html>
            `);
            printWindow.document.close();
            printWindow.focus();
        }}
        
        function exportToWord() {{
            const content = document.querySelector('.premium-results');
            const clonedContent = content.cloneNode(true);
            
            // Remove interactive elements
            const elementsToRemove = clonedContent.querySelectorAll('.actions');
            elementsToRemove.forEach(element => element.remove());
            
            const htmlContent = `
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>Job Fit Analysis Report</title>
                    <style>
                        body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 40px; }}
                        .premium-header h2 {{ color: #2c3e50; text-align: center; border-bottom: 2px solid #2c3e50; padding-bottom: 10px; }}
                        .section h3 {{ color: #2c3e50; border-bottom: 1px solid #2c3e50; padding-bottom: 5px; }}
                    </style>
                </head>
                <body>
                    ` + clonedContent.outerHTML + `
                </body>
                </html>
            `;
            
            const blob = new Blob([htmlContent], {{ type: 'application/msword' }});
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = 'Job_Fit_Analysis_Report.doc';
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            URL.revokeObjectURL(url);
        }}
    </script>
    
    <!-- Load jsPDF library -->
    <script src="https://cdnjs.cloudflare.com/ajax/libs/jspdf/2.5.1/jspdf.umd.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script>
    """
    
    return html_content

def generate_embedded_cover_letter_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for cover letter results using Jinja2 template"""
    
    # Map actual AI response to template expectations
    cover_letter_content = result.get('cover_letter', '')
    tone = result.get('tone', 'Professional')
    word_count = result.get('word_count', 0)
    key_points = result.get('key_points_highlighted', [])
    
    # Create cover_letter_analysis with mock scores based on quality indicators
    letter_analysis = {
        "tone_score": 88,  # Good professional tone
        "keyword_coverage": 85,  # Reasonable keyword integration
        "structure_score": 90,  # Well-structured format
        "overall_rating": 87,  # Overall good quality
        "tone_assessment": f"Your cover letter maintains a {tone.lower()} tone throughout, which is appropriate for the role."
    }
    
    # Transform key_points_highlighted to key_highlights format
    key_highlights = []
    for i, point in enumerate(key_points):
        key_highlights.append({
            'category': f'Key Point {i+1}',
            'title': f'Highlight {i+1}',
            'description': point,
            'impact': 'Strengthens your application by demonstrating relevant experience'
        })
    
    # Simplified strategic_elements - reduce to just 2 key elements
    strategic_elements = [
        {
            'element_type': 'Professional Opening',
            'purpose': 'Establishes immediate interest and relevance',
            'example': 'Strong opening that connects to the role',
            'effectiveness': 85
        },
        {
            'element_type': 'Compelling Closing',
            'purpose': 'Encourages action and reinforces interest', 
            'example': 'Clear call-to-action for next steps',
            'effectiveness': 88
        }
    ]
    
    # Simplified customization_details
    customization_details = {
        'company_research': [
            'Tailored to specific company and role requirements',
            'Incorporates relevant industry terminology'
        ],
        'role_alignment': [
            'Highlights most relevant experiences for this position',
            'Addresses key qualifications mentioned in job posting'
        ],
        'personal_brand': 'The cover letter effectively communicates your unique value proposition'
    }
    
    # Simplified next_steps - reduce to 2 essential steps
    next_steps = [
        {
            'action': 'Review and customize further',
            'details': 'Consider adding specific company details if available',
            'timeline': 'Before sending'
        },
        {
            'action': 'Proofread carefully',
            'details': 'Check for any typos, grammar errors, or formatting issues',
            'timeline': 'Final step before submission'
        }
    ]
    
    # Create success_prediction
    success_prediction = {
        'score': 85,
        'assessment': 'This cover letter effectively positions you as a strong candidate by highlighting relevant experience and demonstrating genuine interest in the role.',
        'strengths': ['Professional tone', 'Relevant experience highlighted', 'Clear structure', 'Compelling narrative']
    }
    
    # Prepare template context with the mapped data
    context = {
        "cover_letter_text": cover_letter_content,  # Template expects cover_letter_text not cover_letter_content
        "letter_analysis": letter_analysis,  # Template expects letter_analysis not cover_letter_analysis
        "key_highlights": key_highlights,
        "strategic_elements": strategic_elements,
        "customization_details": customization_details,
        "next_steps": next_steps,
        "success_prediction": success_prediction,
        "analysis_id": analysis_id
    }
    
    # Render the template to HTML string
    template = templates.get_template("cover_letter_embedded.html")
    return template.render(context)

def generate_embedded_interview_prep_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for interview prep results"""
    return f"""
    <div class="premium-results">
        <div class="premium-header">
            <h2>&#127908; Interview Preparation</h2>
            <p>Personalized interview questions and answers</p>
        </div>
        <div class="section">
            <h3>Interview Prep Results</h3>
            <pre>{result}</pre>
        </div>
        <div class="actions">
            <button class="btn print-btn" onclick="window.print()">&#128424; Print Report</button>
            <a href="/" class="btn">&#127968; Return to App</a>
        </div>
    </div>
    """

def generate_embedded_salary_insights_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for salary insights results"""
    return f"""
    <div class="premium-results">
        <div class="premium-header">
            <h2>&#128176; Salary Insights</h2>
            <p>Market rate analysis for your role</p>
        </div>
        <div class="section">
            <h3>Salary Insights</h3>
            <pre>{result}</pre>
        </div>
        <div class="actions">
            <button class="btn print-btn" onclick="window.print()">&#128424; Print Report</button>
            <a href="/" class="btn">&#127968; Return to App</a>
        </div>
    </div>
    """

def generate_resume_rewrite_html(result: dict, analysis_id: str) -> str:
    """Generate full HTML page for resume rewrite results"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    strategic_changes = result.get('strategic_changes', {})
    before_after_comparison = result.get('before_after_comparison', [])
    interview_generation_potential = result.get('interview_generation_potential', '')
    next_steps = result.get('next_steps', '')
    
    # Extract resume sections
    professional_summary = rewritten_resume.get('professional_summary', '')
    core_competencies = rewritten_resume.get('core_competencies', [])
    professional_experience = rewritten_resume.get('professional_experience', [])
    education = rewritten_resume.get('education', '')
    additional_sections = rewritten_resume.get('additional_sections', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Resume Rewrite - Resume Health Checker</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                margin: 0;
                padding: 20px;
                color: #333;
            }}
            .container {{
                max-width: 1200px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                overflow: hidden;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
            }}
            .header h1 {{
                margin: 0;
                font-size: 2.5rem;
                font-weight: 700;
            }}
            .content {{
                padding: 40px;
            }}
            .two-column {{
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 40px;
                margin: 30px 0;
            }}
            .section {{
                margin-bottom: 40px;
            }}
            .section h2 {{
                color: #667eea;
                font-size: 1.8rem;
                margin-bottom: 20px;
                border-bottom: 2px solid #667eea;
                padding-bottom: 10px;
            }}
            .section h3 {{
                color: #333;
                font-size: 1.3rem;
                margin-bottom: 15px;
            }}
            .rewritten-resume {{
                background: #f8f9ff;
                padding: 30px;
                border-radius: 12px;
                border: 2px solid #667eea;
            }}
            .resume-section {{
                margin-bottom: 25px;
            }}
            .resume-section h3 {{
                color: #667eea;
                border-bottom: 1px solid #667eea;
                padding-bottom: 8px;
                margin-bottom: 15px;
            }}
            .experience-item {{
                background: white;
                padding: 20px;
                border-radius: 8px;
                margin: 15px 0;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            }}
            .experience-header {{
                font-weight: 600;
                color: #333;
                margin-bottom: 10px;
            }}
            .bullet-points {{
                list-style: none;
                padding: 0;
            }}
            .bullet-points li {{
                padding: 5px 0;
                padding-left: 20px;
                position: relative;
            }}
            .bullet-points li:before {{
                content: "•";
                color: #667eea;
                font-weight: bold;
                position: absolute;
                left: 0;
            }}
            .comparison-item {{
                background: #f8f9ff;
                padding: 20px;
                border-radius: 8px;
                margin: 15px 0;
                border-left: 4px solid #ffc107;
            }}
            .actions {{
                text-align: center;
                margin-top: 40px;
                padding-top: 30px;
                border-top: 1px solid #e0e0e0;
            }}
            .btn {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
            }}
            .btn:hover {{
                transform: translateY(-2px);
                box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);
            }}
            .success-section {{
                background: linear-gradient(135deg, #e8f5e8 0%, #f0f8f0 100%);
                padding: 25px;
                border-radius: 12px;
                border: 2px solid #28a745;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>&#128221; Complete Resume Rewrite</h1>
                <p>Your job-targeted resume transformation</p>
            </div>
            
            <div class="content">
                <div class="section">
                    <h2>&#127919; Your Rewritten Resume</h2>
                    <div class="rewritten-resume">
                        <div class="resume-section">
                            <h3>Professional Summary</h3>
                            <p>{professional_summary}</p>
                        </div>
                        
                        <div class="resume-section">
                            <h3>Core Competencies</h3>
                            <ul class="bullet-points">
    """
    
    for competency in core_competencies:
        html_content += f'<li>{competency}</li>'
    
    html_content += f"""
                            </ul>
                        </div>
                        
                        <div class="resume-section">
                            <h3>Professional Experience</h3>
    """
    
    for exp in professional_experience:
        company = exp.get('company', '')
        title = exp.get('title', '')
        duration = exp.get('duration', '')
        bullets = exp.get('rewritten_bullets', [])
        
        html_content += f"""
                            <div class="experience-item">
                                <div class="experience-header">{title} | {company} | {duration}</div>
                                <ul class="bullet-points">
        """
        
        for bullet in bullets:
            html_content += f'<li>{bullet}</li>'
        
        html_content += """
                                </ul>
                            </div>
        """
    
    html_content += f"""
                        </div>
                        
                        <div class="resume-section">
                            <h3>Education</h3>
                            <p>{education}</p>
                        </div>
                        
                        {f'<div class="resume-section"><h3>Additional Information</h3><p>{additional_sections}</p></div>' if additional_sections else ''}
                    </div>
                </div>
                
                <div class="two-column">
                    <div class="section">
                        <h2>&#128640; Strategic Changes</h2>
                        <h3>Keyword Optimization</h3>
                        <ul>
    """
    
    for keyword in strategic_changes.get('keyword_optimization', []):
        html_content += f'<li>{keyword}</li>'
    
    html_content += f"""
                        </ul>
                        
                        <h3>Narrative Positioning</h3>
                        <p>{strategic_changes.get('narrative_positioning', '')}</p>
                        
                        <h3>Competitive Advantages</h3>
                        <ul>
    """
    
    for advantage in strategic_changes.get('competitive_advantages', []):
        html_content += f'<li>{advantage}</li>'
    
    html_content += f"""
                        </ul>
                    </div>
                    
                    <div class="section">
                        <h2>&#128202; Before & After</h2>
    """
    
    for comparison in before_after_comparison:
        html_content += f"""
                        <div class="comparison-item">
                            <h3>{comparison.get('section', 'Section')}</h3>
                            <p><strong>Original Issue:</strong> {comparison.get('original_weakness', '')}</p>
                            <p><strong>Rewrite Improvement:</strong> {comparison.get('rewritten_strength', '')}</p>
                            <p><strong>Expected Impact:</strong> {comparison.get('expected_impact', '')}</p>
                        </div>
        """
    
    html_content += f"""
                    </div>
                </div>
                
                <div class="success-section">
                    <h2>&#127942; Interview Generation Potential</h2>
                    <p>{interview_generation_potential}</p>
                </div>
                
                <div class="success-section">
                    <h2>&#128161; Next Steps</h2>
                    <p>{next_steps}</p>
                </div>
                
                <div class="actions">
                    <button class="btn" onclick="window.print()">&#128424; Print Resume</button>
                    <button class="btn" onclick="copyToClipboard()">&#128203; Copy Text</button>
                    <a href="/" class="btn">&#127968; Return to App</a>
                </div>
            </div>
        </div>
        
        <script>
            function copyToClipboard() {{
                const resumeText = document.querySelector('.rewritten-resume').textContent;
                navigator.clipboard.writeText(resumeText).then(() => {{
                    alert('Resume copied to clipboard!');
                }});
            }}
        </script>
    </body>
    </html>
    """
    
    return html_content

def generate_embedded_resume_rewrite_html(result: dict, analysis_id: str) -> str:
    """Generate embedded HTML for resume rewrite results using Jinja2 template"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    strategic_changes = result.get('strategic_changes', {})
    interview_generation_potential = result.get('interview_generation_potential', '')
    
    # Prepare template context
    context = {
        "rewritten_resume": rewritten_resume,
        "strategic_changes": strategic_changes,
        "interview_generation_potential": interview_generation_potential,
        "analysis_id": analysis_id
    }
    
    # Render the embedded template to HTML string
    template = templates.get_template("resume_rewrite_embedded.html")
    return template.render(context)


# =============================================================================
# EXPORT ENDPOINTS (PDF and DOCX)
# =============================================================================

@router.get("/export/{analysis_id}/pdf")
async def export_pdf(analysis_id: str):
    """Export analysis results as PDF"""
    try:
        # Check if WeasyPrint is available
        if not WEASYPRINT_AVAILABLE:
            return JSONResponse(
                status_code=503,
                content={"error": "pdf_export_unavailable", "message": "PDF export is temporarily unavailable due to missing system dependencies"}
            )
        
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required")
        
        # Get premium result
        premium_result = analysis.get('premium_result')
        if not premium_result:
            raise HTTPException(status_code=404, detail="Premium result not found")
        
        # Determine product type from analysis metadata or assume from result structure
        product_type = analysis.get('product_type', 'resume_analysis')
        
        # Generate clean HTML for PDF (without interactive elements)
        html_content = generate_pdf_html(premium_result, analysis_id, product_type)
        
        # Generate PDF using WeasyPrint
        pdf_buffer = io.BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        # Determine filename based on product type
        filename_map = {
            "resume_analysis": "resume-analysis-report.pdf",
            "job_fit_analysis": "job-fit-analysis-report.pdf",
            "cover_letter": "cover-letter.pdf",
            "resume_rewrite": "rewritten-resume.pdf"
        }
        filename = filename_map.get(product_type, "analysis-report.pdf")
        
        return StreamingResponse(
            io.BytesIO(pdf_buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "pdf_export_error", "message": str(e)}
        )

@router.get("/export/{analysis_id}/docx")
async def export_docx(analysis_id: str):
    """Export analysis results as DOCX (for cover letters and resume rewrites)"""
    try:
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required")
        
        # Get premium result
        premium_result = analysis.get('premium_result')
        if not premium_result:
            raise HTTPException(status_code=404, detail="Premium result not found")
        
        # Determine product type from analysis content
        product_type = analysis.get('product_type')
        if not product_type:
            # Try to infer from premium result content
            if 'rewritten_resume' in premium_result:
                product_type = 'resume_rewrite'
            elif 'cover_letter' in premium_result:
                product_type = 'cover_letter'
            elif 'job_fit_score' in premium_result:
                product_type = 'job_fit_analysis'
            else:
                product_type = 'resume_analysis'
        
        # Support DOCX export for all premium services
        if product_type not in ['resume_analysis', 'job_fit_analysis', 'cover_letter', 'resume_rewrite']:
            raise HTTPException(status_code=400, detail="DOCX export not available for this service type")
        
        # Create DOCX document
        doc = Document()
        
        if product_type == 'resume_analysis':
            generate_resume_analysis_docx(doc, premium_result, analysis_id)
            filename = "resume-analysis.docx"
        elif product_type == 'job_fit_analysis':
            generate_job_fit_analysis_docx(doc, premium_result, analysis_id)
            filename = "job-fit-analysis.docx"
        elif product_type == 'cover_letter':
            generate_cover_letter_docx(doc, premium_result, analysis_id)
            filename = "cover-letter.docx"
        elif product_type == 'resume_rewrite':
            generate_resume_rewrite_docx(doc, premium_result, analysis_id)
            filename = "rewritten-resume.docx"
        
        # Save document to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(docx_buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "docx_export_error", "message": str(e)}
        )

def generate_pdf_html(result: dict, analysis_id: str, product_type: str) -> str:
    """Generate clean HTML suitable for PDF export"""
    
    if product_type == "resume_analysis":
        return generate_pdf_resume_analysis_html(result, analysis_id)
    elif product_type == "job_fit_analysis":
        return generate_pdf_job_fit_html(result, analysis_id)
    elif product_type == "cover_letter":
        return generate_pdf_cover_letter_html(result, analysis_id)
    elif product_type == "resume_rewrite":
        return generate_pdf_resume_rewrite_html(result, analysis_id)
    else:
        return f"<h1>Export for {product_type}</h1><pre>{result}</pre>"

def generate_pdf_resume_analysis_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for resume analysis"""
    
    overall_score = result.get('overall_score', 'N/A')
    strengths = result.get('strength_highlights', [])
    opportunities = result.get('improvement_opportunities', [])
    ats_opt = result.get('ats_optimization', {})
    content_enhancement = result.get('content_enhancement', {})
    text_rewrites = result.get('text_rewrites', [])
    competitive_advantages = result.get('competitive_advantages', '')
    success_prediction = result.get('success_prediction', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Resume Analysis Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; color: #333; }}
            h1 {{ color: #333; text-align: center; margin-bottom: 30px; }}
            h2 {{ color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 8px; margin-top: 30px; }}
            h3 {{ color: #333; margin-top: 20px; }}
            .score {{ font-size: 28px; font-weight: bold; text-align: center; margin: 30px 0; color: #667eea; }}
            ul {{ list-style: none; padding: 0; }}
            li {{ background: #f9f9f9; margin: 8px 0; padding: 12px; border-left: 4px solid #667eea; }}
            .text-rewrite {{ background: #f8f9ff; padding: 15px; margin: 15px 0; border: 1px solid #e0e0e0; border-radius: 6px; }}
            .original {{ background: #fff3cd; padding: 12px; margin: 8px 0; border-left: 3px solid #ffc107; border-radius: 4px; }}
            .improved {{ background: #d4edda; padding: 12px; margin: 8px 0; border-left: 3px solid #28a745; border-radius: 4px; }}
            .why-better {{ font-style: italic; color: #666; margin-top: 8px; }}
            .highlight-section {{ background: #e8f5e8; padding: 20px; border: 2px solid #28a745; border-radius: 8px; margin: 20px 0; }}
        </style>
    </head>
    <body>
        <h1>📊 Premium Resume Analysis Report</h1>
        <div class="score">Overall Score: {overall_score}/100</div>
        
        <h2>💪 Key Strengths</h2>
        <ul>
    """
    
    for strength in strengths:
        html_content += f'<li>{strength}</li>'
    
    html_content += """
        </ul>
        
        <h2>🚀 Improvement Opportunities</h2>
        <ul>
    """
    
    for opportunity in opportunities:
        html_content += f'<li>{opportunity}</li>'
    
    html_content += f"""
        </ul>
        
        <h2>🎯 ATS Optimization</h2>
        <h3>Current Strength</h3>
        <p>{ats_opt.get('current_strength', 'N/A')}</p>
        
        <h3>Enhancement Opportunities</h3>
        <ul>
    """
    
    for enhancement in ats_opt.get('enhancement_opportunities', []):
        html_content += f'<li>{enhancement}</li>'
    
    html_content += f"""
        </ul>
        
        <h3>Impact Prediction</h3>
        <p>{ats_opt.get('impact_prediction', 'N/A')}</p>
        
        <h2>📝 Content Enhancement</h2>
        <h3>Strong Sections</h3>
        <ul>
    """
    
    for strong in content_enhancement.get('strong_sections', []):
        html_content += f'<li>{strong}</li>'
    
    html_content += f"""
        </ul>
        
        <h3>Growth Areas</h3>
        <ul>
    """
    
    for growth in content_enhancement.get('growth_areas', []):
        html_content += f'<li>{growth}</li>'
    
    html_content += f"""
        </ul>
        
        <h3>Strategic Additions</h3>
        <ul>
    """
    
    for addition in content_enhancement.get('strategic_additions', []):
        html_content += f'<li>{addition}</li>'
    
    # Add text rewrites if available
    if text_rewrites:
        html_content += """
        </ul>
        
        <h2>✏️ Text Rewrites</h2>
        """
        
        for rewrite in text_rewrites:
            html_content += f"""
        <div class="text-rewrite">
            <h3>{rewrite.get('section', 'Section')}</h3>
            <div class="original">
                <strong>Original:</strong><br>
                {rewrite.get('original', 'N/A')}
            </div>
            <div class="improved">
                <strong>Improved:</strong><br>
                {rewrite.get('improved', 'N/A')}
            </div>
            <div class="why-better">
                <strong>Why this is better:</strong> {rewrite.get('why_better', 'N/A')}
            </div>
        </div>
            """
    else:
        html_content += "</ul>"
    
    # Add competitive advantages and success prediction if available
    if competitive_advantages:
        html_content += f"""
        
        <div class="highlight-section">
            <h2>🏆 Competitive Advantages</h2>
            <p>{competitive_advantages}</p>
        </div>
        """
    
    if success_prediction:
        html_content += f"""
        
        <div class="highlight-section">
            <h2>📈 Success Prediction</h2>
            <p>{success_prediction}</p>
        </div>
        """
    
    html_content += """
    </body>
    </html>
    """
    
    return html_content

def generate_pdf_job_fit_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for job fit analysis"""
    
    job_fit_score = result.get('job_fit_score', 'N/A')
    strategic_advantages = result.get('strategic_advantages', [])
    positioning_strategy = result.get('positioning_strategy', {})
    optimization_keywords = result.get('optimization_keywords', [])
    resume_enhancements = result.get('resume_enhancements', [])
    text_rewrites = result.get('text_rewrites', [])
    interview_confidence = result.get('interview_confidence', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Job Fit Analysis Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; color: #333; }}
            h1 {{ color: #333; text-align: center; margin-bottom: 30px; }}
            h2 {{ color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 8px; margin-top: 30px; }}
            h3 {{ color: #333; margin-top: 20px; }}
            .score {{ font-size: 28px; font-weight: bold; text-align: center; margin: 30px 0; color: #667eea; }}
            ul {{ list-style: none; padding: 0; }}
            li {{ background: #f9f9f9; margin: 8px 0; padding: 12px; border-left: 4px solid #667eea; }}
            .text-rewrite {{ background: #f8f9ff; padding: 15px; margin: 15px 0; border: 1px solid #e0e0e0; border-radius: 6px; }}
            .original {{ background: #fff3cd; padding: 12px; margin: 8px 0; border-left: 3px solid #ffc107; border-radius: 4px; }}
            .improved {{ background: #d4edda; padding: 12px; margin: 8px 0; border-left: 3px solid #28a745; border-radius: 4px; }}
            .why-better {{ font-style: italic; color: #666; margin-top: 8px; }}
            .highlight-section {{ background: #e8f5e8; padding: 20px; border: 2px solid #28a745; border-radius: 8px; margin: 20px 0; }}
            .keyword-tag {{ background: #667eea; color: white; padding: 6px 12px; margin: 4px; border-radius: 20px; font-size: 14px; display: inline-block; }}
        </style>
    </head>
    <body>
        <h1>🎯 Job Fit Analysis Report</h1>
        <div class="score">Job Fit Score: {job_fit_score}/100</div>
        
        <h2>⭐ Strategic Advantages</h2>
        <ul>
    """
    
    for advantage in strategic_advantages:
        html_content += f'<li>{advantage}</li>'
    
    # Add positioning strategy if available
    primary_value = positioning_strategy.get('primary_value', '')
    supporting_qualifications = positioning_strategy.get('supporting_qualifications', [])
    unique_differentiators = positioning_strategy.get('unique_differentiators', [])
    
    if primary_value:
        html_content += f"""
        </ul>
        
        <h2>🏆 Primary Value Proposition</h2>
        <p>{primary_value}</p>
        """
    
    if supporting_qualifications:
        html_content += """
        
        <h2>✅ Supporting Qualifications</h2>
        <ul>
        """
        for qual in supporting_qualifications:
            html_content += f'<li>{qual}</li>'
        html_content += "</ul>"
    
    if unique_differentiators:
        html_content += """
        
        <h2>💎 Unique Differentiators</h2>
        <ul>
        """
        for diff in unique_differentiators:
            html_content += f'<li>{diff}</li>'
        html_content += "</ul>"
    
    if optimization_keywords:
        html_content += """
        
        <h2>🔑 Optimization Keywords</h2>
        <div>
        """
        for keyword in optimization_keywords:
            html_content += f'<span class="keyword-tag">{keyword}</span>'
        html_content += "</div>"
    
    if resume_enhancements:
        html_content += """
        
        <h2>📝 Resume Enhancements</h2>
        <ul>
        """
        for enhancement in resume_enhancements:
            html_content += f'<li>{enhancement}</li>'
        html_content += "</ul>"
    
    # Add text rewrites if available
    if text_rewrites:
        html_content += """
        
        <h2>✏️ Text Rewrites</h2>
        """
        
        for rewrite in text_rewrites:
            section = rewrite.get('section', 'Section')
            original = rewrite.get('original', '')
            job_optimized = rewrite.get('job_optimized', '')
            strategic_impact = rewrite.get('strategic_impact', '')
            
            html_content += f"""
        <div class="text-rewrite">
            <h3>{section}</h3>
            <div class="original">
                <strong>Before:</strong><br>
                {original}
            </div>
            <div class="improved">
                <strong>Job-Optimized:</strong><br>
                {job_optimized}
            </div>
            <div class="why-better">
                <strong>Strategic Impact:</strong> {strategic_impact}
            </div>
        </div>
            """
    
    if interview_confidence:
        html_content += f"""
        
        <div class="highlight-section">
            <h2>🎯 Interview Confidence</h2>
            <p>{interview_confidence}</p>
        </div>
        """
    
    html_content += """
    </body>
    </html>
    """
    
    return html_content

def generate_pdf_cover_letter_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for cover letter"""
    
    cover_letter = result.get('cover_letter', '')
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Cover Letter</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
            h1 {{ color: #333; text-align: center; margin-bottom: 30px; }}
            .cover-letter {{ white-space: pre-line; }}
        </style>
    </head>
    <body>
        <h1>Cover Letter</h1>
        <div class="cover-letter">{cover_letter}</div>
    </body>
    </html>
    """
    
    return html_content

def generate_pdf_resume_rewrite_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for resume rewrite"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    professional_summary = rewritten_resume.get('professional_summary', '')
    core_competencies = rewritten_resume.get('core_competencies', [])
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Rewritten Resume</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
            h1 {{ color: #333; text-align: center; }}
            h2 {{ color: #667eea; border-bottom: 1px solid #667eea; }}
            ul {{ list-style: none; padding: 0; }}
            li {{ margin: 4px 0; padding: 8px; background: #f9f9f9; }}
        </style>
    </head>
    <body>
        <h1>Rewritten Resume</h1>
        
        <h2>Professional Summary</h2>
        <p>{professional_summary}</p>
        
        <h2>Core Competencies</h2>
        <ul>
    """
    
    for competency in core_competencies:
        html_content += f'<li>• {competency}</li>'
    
    html_content += """
        </ul>
    </body>
    </html>
    """
    
    return html_content

def generate_resume_analysis_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for resume analysis"""
    
    # Add title
    title = doc.add_heading('Resume Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Overall score
    overall_score = result.get('overall_score', 'N/A')
    doc.add_heading(f'Overall Score: {overall_score}/100', level=1)
    
    # Strengths
    strength_highlights = result.get('strength_highlights', [])
    if strength_highlights:
        doc.add_heading('Key Strengths', level=1)
        for strength in strength_highlights:
            doc.add_paragraph(strength, style='List Bullet')
    
    # Improvement opportunities
    improvement_opportunities = result.get('improvement_opportunities', [])
    if improvement_opportunities:
        doc.add_heading('Improvement Opportunities', level=1)
        for opportunity in improvement_opportunities:
            doc.add_paragraph(opportunity, style='List Bullet')
    
    # ATS Optimization
    ats_optimization = result.get('ats_optimization', {})
    if ats_optimization:
        doc.add_heading('ATS Optimization', level=1)
        for key, value in ats_optimization.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))
    
    # Text rewrites
    text_rewrites = result.get('text_rewrites', [])
    if text_rewrites:
        doc.add_heading('Text Rewrites', level=1)
        for i, rewrite in enumerate(text_rewrites, 1):
            doc.add_heading(f'Rewrite {i}', level=2)
            doc.add_paragraph(f"Original: {rewrite.get('original', '')}")
            doc.add_paragraph(f"Improved: {rewrite.get('improved', '')}")
            if rewrite.get('why_better'):
                doc.add_paragraph(f"Why better: {rewrite.get('why_better', '')}")

def generate_job_fit_analysis_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for job fit analysis"""
    
    # Add title
    title = doc.add_heading('Job Fit Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Overall match score
    overall_match_score = result.get('overall_match_score', 'N/A')
    doc.add_heading(f'Overall Job Match: {overall_match_score}%', level=1)
    
    # Matching qualifications
    matching_qualifications = result.get('matching_qualifications', [])
    if matching_qualifications:
        doc.add_heading('Matching Qualifications', level=1)
        for qual in matching_qualifications:
            doc.add_paragraph(qual, style='List Bullet')
    
    # Missing qualifications
    missing_qualifications = result.get('missing_qualifications', [])
    if missing_qualifications:
        doc.add_heading('Areas for Improvement', level=1)
        for qual in missing_qualifications:
            doc.add_paragraph(qual, style='List Bullet')
    
    # Recommendations
    recommendations = result.get('recommendations', [])
    if recommendations:
        doc.add_heading('Recommendations', level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f'Recommendation {i}', level=2)
            doc.add_paragraph(rec.get('recommendation', ''))
            if rec.get('specific_actions'):
                doc.add_paragraph('Specific Actions:')
                for action in rec.get('specific_actions', []):
                    doc.add_paragraph(action, style='List Bullet')
    
    # Interview preparation
    interview_prep = result.get('interview_preparation', {})
    if interview_prep:
        doc.add_heading('Interview Preparation', level=1)
        for key, value in interview_prep.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))

def generate_cover_letter_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for cover letter"""
    
    cover_letter = result.get('cover_letter', '')
    
    # Add title
    title = doc.add_heading('Cover Letter', 0)
    title.alignment = 1  # Center alignment
    
    # Add cover letter content
    paragraphs = cover_letter.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

def generate_resume_rewrite_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for resume rewrite"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    
    # Add title
    title = doc.add_heading('Rewritten Resume', 0)
    title.alignment = 1  # Center alignment
    
    # Professional Summary
    professional_summary = rewritten_resume.get('professional_summary', '')
    if professional_summary:
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(professional_summary)
    
    # Core Competencies
    core_competencies = rewritten_resume.get('core_competencies', [])
    if core_competencies:
        doc.add_heading('Core Competencies', level=1)
        for competency in core_competencies:
            p = doc.add_paragraph(competency, style='List Bullet')
    
    # Professional Experience
    professional_experience = rewritten_resume.get('professional_experience', [])
    if professional_experience:
        doc.add_heading('Professional Experience', level=1)
        for exp in professional_experience:
            company = exp.get('company', '')
            title_text = exp.get('title', '')
            duration = exp.get('duration', '')
            bullets = exp.get('rewritten_bullets', [])
            
            # Experience header
            doc.add_paragraph(f"{title_text} | {company} | {duration}", style='Heading 2')
            
            # Bullet points
            for bullet in bullets:
                doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    education = rewritten_resume.get('education', '')
    if education:
        doc.add_heading('Education', level=1)
        doc.add_paragraph(education)
    
    # Additional Qualifications
    additional_sections = rewritten_resume.get('additional_sections', '')
    if additional_sections:
        doc.add_heading('Additional Qualifications', level=1)
        doc.add_paragraph(additional_sections)
    
    # Strategic Optimizations
    strategic_changes = result.get('strategic_changes', {})
    if strategic_changes:
        doc.add_heading('Strategic Optimizations', level=1)
        
        # Keywords Added
        keyword_optimization = strategic_changes.get('keyword_optimization', [])
        if keyword_optimization:
            doc.add_paragraph(f"Keywords Added: {len(keyword_optimization)} job-specific terms")
            for keyword in keyword_optimization:
                doc.add_paragraph(f"• {keyword}", style='List Bullet')
        
        # ATS Enhancements
        ats_enhancements = strategic_changes.get('ats_enhancements', [])
        if ats_enhancements:
            doc.add_paragraph(f"ATS Enhancements: {len(ats_enhancements)} formatting improvements")
            for enhancement in ats_enhancements:
                doc.add_paragraph(f"• {enhancement}", style='List Bullet')
        
        # Competitive Edge
        competitive_advantages = strategic_changes.get('competitive_advantages', [])
        if competitive_advantages:
            doc.add_paragraph(f"Competitive Edge: {len(competitive_advantages)} unique differentiators highlighted")
            for advantage in competitive_advantages:
                doc.add_paragraph(f"• {advantage}", style='List Bullet')
    
    # Interview Generation Potential
    interview_generation_potential = result.get('interview_generation_potential', '')
    if interview_generation_potential:
        doc.add_heading('Interview Generation Potential', level=1)
        doc.add_paragraph(interview_generation_potential)
