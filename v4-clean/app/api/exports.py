"""
Export Routes for Resume Health Checker v4.0

All export-related endpoints including PDF and DOCX generation
with dedicated HTML generation functions for premium results.
"""
import logging
import io
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse
from fastapi.templating import Jinja2Templates

# Try to import WeasyPrint, fall back to alternative approach if not available
try:
    # WeasyPrint requires system dependencies (gobject-2.0, pango, etc.)
    # Commented out until system dependencies are installed
    # from weasyprint import HTML, CSS
    HTML = None
    CSS = None
    WEASYPRINT_AVAILABLE = False
except ImportError:
    HTML = None
    CSS = None
    WEASYPRINT_AVAILABLE = False

from docx import Document

from ..core.database import AnalysisDB

logger = logging.getLogger(__name__)

# Create router
router = APIRouter()

# Setup Jinja2 templates
templates = Jinja2Templates(directory="app/templates")

@router.get("/export/{analysis_id}/pdf")
async def export_pdf(analysis_id: str):
    """Export analysis results as PDF"""
    try:
        # For now, provide HTML version instead of PDF when WeasyPrint unavailable
        if not WEASYPRINT_AVAILABLE:
            logger.warning(f"PDF export requested for {analysis_id} but WeasyPrint not available, returning HTML")
            # Get analysis data
            analysis = AnalysisDB.get(analysis_id)
            if not analysis:
                raise HTTPException(status_code=404, detail="Analysis not found")
            
            # Check if payment was successful
            if analysis.get('payment_status') != 'paid':
                raise HTTPException(status_code=402, detail="Payment required")
            
            # Get premium result
            premium_result = analysis.get('premium_result')
            if not premium_result:
                raise HTTPException(status_code=404, detail="Premium result not found")
            
            # Determine product type from analysis metadata or assume from result structure
            product_type = analysis.get('product_type', 'resume_analysis')
            
            # Generate clean HTML for PDF export (without interactive elements)
            html_content = generate_pdf_html(premium_result, analysis_id, product_type)
            
            # Return HTML with PDF-like styling as fallback
            return HTMLResponse(
                content=html_content,
                headers={"Content-Disposition": f"attachment; filename=analysis-{analysis_id}.html"}
            )
        
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required")
        
        # Get premium result
        premium_result = analysis.get('premium_result')
        if not premium_result:
            raise HTTPException(status_code=404, detail="Premium result not found")
        
        # Determine product type from analysis metadata or assume from result structure
        product_type = analysis.get('product_type', 'resume_analysis')
        
        # Generate clean HTML for PDF (without interactive elements)
        html_content = generate_pdf_html(premium_result, analysis_id, product_type)
        
        # Generate PDF using WeasyPrint
        pdf_buffer = io.BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        # Determine filename based on product type
        filename_map = {
            "resume_analysis": "resume-analysis-report.pdf",
            "job_fit_analysis": "job-fit-analysis-report.pdf",
            "cover_letter": "cover-letter.pdf",
            "resume_rewrite": "rewritten-resume.pdf"
        }
        filename = filename_map.get(product_type, "analysis-report.pdf")
        
        return StreamingResponse(
            io.BytesIO(pdf_buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "pdf_export_error", "message": str(e)}
        )

@router.get("/export/{analysis_id}/docx")
async def export_docx(analysis_id: str):
    """Export analysis results as DOCX (for cover letters and resume rewrites)"""
    try:
        # Get analysis data
        analysis = AnalysisDB.get(analysis_id)
        if not analysis:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        # Check if payment was successful
        if analysis.get('payment_status') != 'paid':
            raise HTTPException(status_code=402, detail="Payment required")
        
        # Get premium result
        premium_result = analysis.get('premium_result')
        if not premium_result:
            raise HTTPException(status_code=404, detail="Premium result not found")
        
        # Determine product type from analysis content
        product_type = analysis.get('product_type')
        if not product_type:
            # Try to infer from premium result content
            if 'rewritten_resume' in premium_result:
                product_type = 'resume_rewrite'
            elif 'cover_letter' in premium_result:
                product_type = 'cover_letter'
            elif 'job_fit_score' in premium_result:
                product_type = 'job_fit_analysis'
            else:
                product_type = 'resume_analysis'
        
        # Support DOCX export for all premium services
        if product_type not in ['resume_analysis', 'job_fit_analysis', 'cover_letter', 'resume_rewrite']:
            raise HTTPException(status_code=400, detail="DOCX export not available for this service type")
        
        # Create DOCX document
        doc = Document()
        
        if product_type == 'resume_analysis':
            generate_resume_analysis_docx(doc, premium_result, analysis_id)
            filename = "resume-analysis.docx"
        elif product_type == 'job_fit_analysis':
            generate_job_fit_analysis_docx(doc, premium_result, analysis_id)
            filename = "job-fit-analysis.docx"
        elif product_type == 'cover_letter':
            generate_cover_letter_docx(doc, premium_result, analysis_id)
            filename = "cover-letter.docx"
        elif product_type == 'resume_rewrite':
            generate_resume_rewrite_docx(doc, premium_result, analysis_id)
            filename = "rewritten-resume.docx"
        
        # Save document to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(docx_buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": "docx_export_error", "message": str(e)}
        )

# =============================================================================
# PDF HTML GENERATION FUNCTIONS
# =============================================================================

def generate_pdf_html(result: dict, analysis_id: str, product_type: str) -> str:
    """Generate clean HTML suitable for PDF export"""
    
    if product_type == "resume_analysis":
        return generate_pdf_resume_analysis_html(result, analysis_id)
    elif product_type == "job_fit_analysis":
        return generate_pdf_job_fit_html(result, analysis_id)
    elif product_type == "cover_letter":
        return generate_pdf_cover_letter_html(result, analysis_id)
    elif product_type == "resume_rewrite":
        return generate_pdf_resume_rewrite_html(result, analysis_id)
    else:
        return f"<h1>Export for {product_type}</h1><pre>{result}</pre>"

def generate_pdf_resume_analysis_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for resume analysis using template"""
    
    context = {
        "overall_score": result.get('overall_score') or 'N/A',
        "strengths": result.get('strength_highlights') or [],
        "opportunities": result.get('improvement_opportunities') or [],
        "ats_opt": result.get('ats_optimization') or {},
        "content_enhancement": result.get('content_enhancement') or {},
        "text_rewrites": result.get('text_rewrites') or [],
        "competitive_advantages": result.get('competitive_advantages') or '',
        "success_prediction": result.get('success_prediction') or ''
    }
    
    template = templates.get_template("pdf_resume_analysis.html")
    return template.render(context)

def generate_pdf_job_fit_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for job fit analysis using template"""
    
    context = {
        "job_fit_score": result.get('job_fit_score') or 'N/A',
        "strategic_advantages": result.get('strategic_advantages') or [],
        "positioning_strategy": result.get('positioning_strategy') or {},
        "optimization_keywords": result.get('optimization_keywords') or [],
        "resume_enhancements": result.get('resume_enhancements') or [],
        "text_rewrites": result.get('text_rewrites') or [],
        "interview_confidence": result.get('interview_confidence') or ''
    }
    
    template = templates.get_template("pdf_job_fit_analysis.html")
    return template.render(context)

def generate_pdf_cover_letter_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for cover letter using template"""
    
    context = {
        "cover_letter": result.get('cover_letter') or ''
    }
    
    template = templates.get_template("pdf_cover_letter.html")
    return template.render(context)

def generate_pdf_resume_rewrite_html(result: dict, analysis_id: str) -> str:
    """Generate PDF-optimized HTML for resume rewrite using template"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    
    context = {
        "professional_summary": rewritten_resume.get('professional_summary') or '',
        "core_competencies": rewritten_resume.get('core_competencies') or []
    }
    
    template = templates.get_template("pdf_resume_rewrite.html")
    return template.render(context)

# =============================================================================
# DOCX GENERATION FUNCTIONS
# =============================================================================

def generate_resume_analysis_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for resume analysis"""
    
    # Add title
    title = doc.add_heading('Resume Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Overall score
    overall_score = result.get('overall_score', 'N/A')
    doc.add_heading(f'Overall Score: {overall_score}/100', level=1)
    
    # Strengths
    strength_highlights = result.get('strength_highlights', [])
    if strength_highlights:
        doc.add_heading('Key Strengths', level=1)
        for strength in strength_highlights:
            doc.add_paragraph(strength, style='List Bullet')
    
    # Improvement opportunities
    improvement_opportunities = result.get('improvement_opportunities', [])
    if improvement_opportunities:
        doc.add_heading('Improvement Opportunities', level=1)
        for opportunity in improvement_opportunities:
            doc.add_paragraph(opportunity, style='List Bullet')
    
    # ATS Optimization
    ats_optimization = result.get('ats_optimization', {})
    if ats_optimization:
        doc.add_heading('ATS Optimization', level=1)
        for key, value in ats_optimization.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))
    
    # Text rewrites
    text_rewrites = result.get('text_rewrites', [])
    if text_rewrites:
        doc.add_heading('Text Rewrites', level=1)
        for i, rewrite in enumerate(text_rewrites, 1):
            doc.add_heading(f'Rewrite {i}', level=2)
            doc.add_paragraph(f"Original: {rewrite.get('original', '')}")
            doc.add_paragraph(f"Improved: {rewrite.get('improved', '')}")
            if rewrite.get('why_better'):
                doc.add_paragraph(f"Why better: {rewrite.get('why_better', '')}")

def generate_job_fit_analysis_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for job fit analysis"""
    
    # Add title
    title = doc.add_heading('Job Fit Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Overall match score
    overall_match_score = result.get('overall_match_score', 'N/A')
    doc.add_heading(f'Overall Job Match: {overall_match_score}%', level=1)
    
    # Matching qualifications
    matching_qualifications = result.get('matching_qualifications', [])
    if matching_qualifications:
        doc.add_heading('Matching Qualifications', level=1)
        for qual in matching_qualifications:
            doc.add_paragraph(qual, style='List Bullet')
    
    # Missing qualifications
    missing_qualifications = result.get('missing_qualifications', [])
    if missing_qualifications:
        doc.add_heading('Areas for Improvement', level=1)
        for qual in missing_qualifications:
            doc.add_paragraph(qual, style='List Bullet')
    
    # Recommendations
    recommendations = result.get('recommendations', [])
    if recommendations:
        doc.add_heading('Recommendations', level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f'Recommendation {i}', level=2)
            doc.add_paragraph(rec.get('recommendation', ''))
            if rec.get('specific_actions'):
                doc.add_paragraph('Specific Actions:')
                for action in rec.get('specific_actions', []):
                    doc.add_paragraph(action, style='List Bullet')
    
    # Interview preparation
    interview_prep = result.get('interview_preparation', {})
    if interview_prep:
        doc.add_heading('Interview Preparation', level=1)
        for key, value in interview_prep.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))

def generate_cover_letter_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for cover letter"""
    
    cover_letter = result.get('cover_letter', '')
    
    # Add title
    title = doc.add_heading('Cover Letter', 0)
    title.alignment = 1  # Center alignment
    
    # Add cover letter content
    paragraphs = cover_letter.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

def generate_resume_rewrite_docx(doc: Document, result: dict, analysis_id: str) -> None:
    """Generate DOCX document for resume rewrite"""
    
    rewritten_resume = result.get('rewritten_resume', {})
    
    # Add title
    title = doc.add_heading('Rewritten Resume', 0)
    title.alignment = 1  # Center alignment
    
    # Professional Summary
    professional_summary = rewritten_resume.get('professional_summary', '')
    if professional_summary:
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(professional_summary)
    
    # Core Competencies
    core_competencies = rewritten_resume.get('core_competencies', [])
    if core_competencies:
        doc.add_heading('Core Competencies', level=1)
        for competency in core_competencies:
            p = doc.add_paragraph(competency, style='List Bullet')
    
    # Professional Experience
    professional_experience = rewritten_resume.get('professional_experience', [])
    if professional_experience:
        doc.add_heading('Professional Experience', level=1)
        for exp in professional_experience:
            company = exp.get('company', '')
            title_text = exp.get('title', '')
            duration = exp.get('duration', '')
            bullets = exp.get('rewritten_bullets', [])
            
            # Experience header
            doc.add_paragraph(f"{title_text} | {company} | {duration}", style='Heading 2')
            
            # Bullet points
            for bullet in bullets:
                doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    education = rewritten_resume.get('education', '')
    if education:
        doc.add_heading('Education', level=1)
        doc.add_paragraph(education)
    
    # Additional Qualifications
    additional_sections = rewritten_resume.get('additional_sections', '')
    if additional_sections:
        doc.add_heading('Additional Qualifications', level=1)
        doc.add_paragraph(additional_sections)
    
    # Note: Strategic Optimizations and Interview Generation Potential are analysis content
    # and should NOT be included in the resume document export