import pytest
import tempfile
import os
from fastapi.testclient import TestClient
from unittest.mock import Mock, patch
import json
from docx import Document
import fitz
import io

# Import our app
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from main import app

@pytest.fixture
def client():
    """FastAPI test client"""
    return TestClient(app)

@pytest.fixture
def mock_openai_response():
    """Mock OpenAI API response"""
    return {
        "choices": [
            {
                "message": {
                    "content": json.dumps({
                        "overall_score": 75,
                        "major_issues": [
                            "Missing quantifiable achievements",
                            "Generic job descriptions", 
                            "Poor formatting consistency"
                        ],
                        "teaser_message": "Get detailed feedback to boost your interview rate by 3x!"
                    })
                }
            }
        ]
    }

@pytest.fixture
def mock_openai_paid_response():
    """Mock OpenAI API response for paid analysis"""
    return {
        "choices": [
            {
                "message": {
                    "content": json.dumps({
                        "overall_score": 75,
                        "ats_optimization": {
                            "score": 70,
                            "issues": ["Missing keywords", "Complex formatting"],
                            "improvements": ["Add industry keywords", "Simplify layout"]
                        },
                        "content_clarity": {
                            "score": 80,
                            "issues": ["Vague accomplishments"],
                            "improvements": ["Add specific metrics"]
                        },
                        "impact_metrics": {
                            "score": 60,
                            "issues": ["No quantified results"],
                            "improvements": ["Add percentage improvements"]
                        },
                        "formatting": {
                            "score": 85,
                            "issues": ["Inconsistent spacing"],
                            "improvements": ["Standardize formatting"]
                        },
                        "top_recommendations": [
                            "Add quantifiable achievements",
                            "Include industry keywords",
                            "Improve formatting consistency"
                        ]
                    })
                }
            }
        ]
    }

@pytest.fixture
def sample_pdf_file():
    """Create a sample PDF file for testing"""
    # Create a simple PDF with sample resume content
    pdf_content = b"""Sample Resume Content
John Doe
Software Developer
Experience: 5 years developing web applications
Skills: Python, JavaScript, React"""
    
    # Create a temporary PDF file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    
    # Create a simple PDF using PyMuPDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Sample Resume Content\n\nJohn Doe\nSoftware Developer\n\nExperience: 5 years developing web applications\nSkills: Python, JavaScript, React")
    doc.save(temp_file.name)
    doc.close()
    
    with open(temp_file.name, 'rb') as f:
        content = f.read()
    
    os.unlink(temp_file.name)
    return content

@pytest.fixture
def sample_docx_file():
    """Create a sample DOCX file for testing"""
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_heading('Software Developer', level=1)
    doc.add_paragraph('Experience: 5 years developing web applications')
    doc.add_paragraph('Skills: Python, JavaScript, React')
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

@pytest.fixture
def invalid_file_content():
    """Create invalid file content for testing"""
    return b"This is not a valid PDF or DOCX file"

@pytest.fixture
def mock_env_vars():
    """Mock environment variables"""
    with patch.dict(os.environ, {
        'OPENAI_API_KEY': 'test-api-key',
        'STRIPE_PAYMENT_SUCCESS_TOKEN': 'test-payment-token'
    }):
        yield