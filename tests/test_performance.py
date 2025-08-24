import pytest
import time
import asyncio
from unittest.mock import patch
import concurrent.futures


class TestPerformanceMetrics:
    """Performance and load testing"""
    
    @patch('main.get_ai_analysis')
    def test_response_time_under_load(self, mock_ai_analysis, client, sample_pdf_file):
        """Test response times under concurrent load"""
        mock_ai_analysis.return_value = {
            "overall_score": 75,
            "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
            "teaser_message": "Test message"
        }
        
        def make_request():
            start_time = time.time()
            response = client.post(
                "/api/check-resume",
                files={"file": ("resume.pdf", sample_pdf_file, "application/pdf")}
            )
            end_time = time.time()
            return response.status_code, end_time - start_time
        
        # Test concurrent requests
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(make_request) for _ in range(10)]
            results = [future.result() for future in concurrent.futures.as_completed(futures)]
        
        # All requests should succeed
        status_codes = [result[0] for result in results]
        response_times = [result[1] for result in results]
        
        assert all(status == 200 for status in status_codes)
        # 95th percentile should be under 5 seconds
        response_times.sort()
        p95_time = response_times[int(0.95 * len(response_times))]
        assert p95_time < 5.0, f"95th percentile response time {p95_time}s exceeds 5s threshold"
    
    @patch('main.get_ai_analysis')  
    def test_memory_usage_stability(self, mock_ai_analysis, client, sample_pdf_file):
        """Test that memory usage remains stable across multiple requests"""
        import psutil
        import os
        
        mock_ai_analysis.return_value = {
            "overall_score": 75,
            "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
            "teaser_message": "Test message"
        }
        
        process = psutil.Process(os.getpid())
        initial_memory = process.memory_info().rss / 1024 / 1024  # MB
        
        # Make 20 requests
        for i in range(20):
            response = client.post(
                "/api/check-resume", 
                files={"file": ("resume.pdf", sample_pdf_file, "application/pdf")}
            )
            assert response.status_code == 200
            
            # Check memory every 5 requests
            if i % 5 == 0:
                current_memory = process.memory_info().rss / 1024 / 1024
                memory_increase = current_memory - initial_memory
                # Memory shouldn't increase by more than 100MB
                assert memory_increase < 100, f"Memory increased by {memory_increase}MB"
    
    def test_file_processing_performance(self, sample_pdf_file, sample_docx_file):
        """Test file processing performance for different file types"""
        from main import extract_text_from_pdf, extract_text_from_docx
        
        # Test PDF processing time
        start_time = time.time()
        with patch('main.fitz.open') as mock_fitz:
            mock_doc = type('MockDoc', (), {
                '__iter__': lambda self: iter([type('MockPage', (), {'get_text': lambda: 'test text'})()]),
                'close': lambda self: None
            })()
            mock_fitz.return_value = mock_doc
            
            pdf_text = extract_text_from_pdf(sample_pdf_file)
            pdf_time = time.time() - start_time
        
        # Test DOCX processing time
        start_time = time.time()
        docx_text = extract_text_from_docx(sample_docx_file)
        docx_time = time.time() - start_time
        
        # Both should process within reasonable time
        assert pdf_time < 2.0, f"PDF processing took {pdf_time}s, should be under 2s"
        assert docx_time < 2.0, f"DOCX processing took {docx_time}s, should be under 2s"
        
        # Results should not be empty
        assert len(pdf_text) > 0
        assert len(docx_text) > 0


class TestScalabilityConsiderations:
    """Test scalability limits and considerations"""
    
    def test_large_file_handling(self, client):
        """Test handling of large resume files"""
        # Create a large but valid DOCX content
        from docx import Document
        import io
        
        doc = Document()
        doc.add_heading('Large Resume Test', 0)
        
        # Add lots of content to simulate a large resume
        for i in range(100):
            doc.add_paragraph(f'Experience entry {i}: ' + 'This is detailed experience description. ' * 20)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        large_docx_content = file_stream.read()
        
        # File should be several MB
        assert len(large_docx_content) > 1000000  # > 1MB
        
        with patch('main.get_ai_analysis') as mock_ai:
            mock_ai.return_value = {
                "overall_score": 75,
                "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
                "teaser_message": "Test message"
            }
            
            start_time = time.time()
            response = client.post(
                "/api/check-resume",
                files={"file": ("large_resume.docx", large_docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
            processing_time = time.time() - start_time
        
        # Should handle large files gracefully
        assert response.status_code == 200
        # Should not take too long even for large files
        assert processing_time < 10.0, f"Large file processing took {processing_time}s"
    
    @patch('main.get_ai_analysis')
    def test_rate_limiting_consideration(self, mock_ai_analysis, client, sample_pdf_file):
        """Test rapid successive requests (rate limiting consideration)"""
        mock_ai_analysis.return_value = {
            "overall_score": 75,
            "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
            "teaser_message": "Test message"
        }
        
        # Make rapid successive requests
        responses = []
        for i in range(10):
            response = client.post(
                "/api/check-resume",
                files={"file": ("resume.pdf", sample_pdf_file, "application/pdf")}
            )
            responses.append(response)
            time.sleep(0.1)  # Small delay
        
        # All should succeed (no rate limiting implemented yet)
        assert all(r.status_code == 200 for r in responses)
        
        # Note: In production, you might want to implement rate limiting
        # This test documents current behavior and can be updated when rate limiting is added


class TestResourceUsage:
    """Test resource usage patterns"""
    
    @patch('main.get_ai_analysis')
    def test_cpu_usage_reasonable(self, mock_ai_analysis, client, sample_pdf_file):
        """Test that CPU usage remains reasonable"""
        import psutil
        
        mock_ai_analysis.return_value = {
            "overall_score": 75,
            "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
            "teaser_message": "Test message"
        }
        
        # Monitor CPU usage
        cpu_percent_before = psutil.cpu_percent(interval=1)
        
        # Make several requests
        for _ in range(5):
            response = client.post(
                "/api/check-resume",
                files={"file": ("resume.pdf", sample_pdf_file, "application/pdf")}
            )
            assert response.status_code == 200
        
        cpu_percent_after = psutil.cpu_percent(interval=1)
        
        # CPU usage shouldn't spike dramatically
        # (This is a rough check, actual values depend on system)
        cpu_increase = cpu_percent_after - cpu_percent_before
        assert cpu_increase < 50, f"CPU usage increased by {cpu_increase}%"
    
    def test_temporary_file_cleanup(self, client, sample_pdf_file):
        """Test that temporary files are properly cleaned up"""
        import tempfile
        import os
        
        temp_dir = tempfile.gettempdir()
        temp_files_before = len([f for f in os.listdir(temp_dir) if f.startswith('tmp')])
        
        with patch('main.get_ai_analysis') as mock_ai:
            mock_ai.return_value = {
                "overall_score": 75,
                "major_issues": ["Issue 1", "Issue 2", "Issue 3"],
                "teaser_message": "Test message"
            }
            
            # Make several requests that process files
            for _ in range(3):
                response = client.post(
                    "/api/check-resume",
                    files={"file": ("resume.pdf", sample_pdf_file, "application/pdf")}
                )
                assert response.status_code == 200
        
        # Allow some time for cleanup
        time.sleep(1)
        
        temp_files_after = len([f for f in os.listdir(temp_dir) if f.startswith('tmp')])
        
        # Temporary files should be cleaned up (or at least not accumulate significantly)
        temp_file_increase = temp_files_after - temp_files_before
        assert temp_file_increase <= 1, f"Temporary files increased by {temp_file_increase}"